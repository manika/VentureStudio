"""
PDF document generator for Venture Studio AI Advisor.
Generates professional regulatory documents (Quality Manual, SOPs, etc.)
formatted to match Imbed Biosciences document control standards.
"""
import io
import re
from datetime import date
from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfgen.canvas import Canvas
from reportlab.platypus import (
    HRFlowable,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from config import OUTPUTS_DIR, MAX_OUTPUT_TOKENS
from modules.llm_client import safe_generate

_IMBED_BASE = (
    "/Users/manika/My Drive/VentureStudio/data"
    "/parent_company_raw/Knowledgebase/QualitySystem/All Files to Inspector 2"
)
_D1 = f"{_IMBED_BASE}/Day 1 - 10FEB2025"
_D2 = f"{_IMBED_BASE}/Day 2 - 11FEB2025"
_D3 = f"{_IMBED_BASE}/Day 3 - 12FEB2025"
_D4 = f"{_IMBED_BASE}/Day 4 - 13FEB2025"

# Legacy alias kept so app.py import doesn't break
IMBED_QM_PATH = f"{_D1}/7 42002_Imbed Quality Manual.pdf"

# Maps each document type to the actual Imbed reference SOP/document
DOC_REFERENCE_PATHS: dict[str, str] = {
    "Quality Manual":                                      f"{_D1}/7 42002_Imbed Quality Manual.pdf",
    "Document and Change Control":                         f"{_D1}/18 43001 SOP Document and Change Control.pdf",
    "Corrective and Preventive Action (CAPA)":             f"{_D2}/26 43013 SOP Corrective and Preventive Action - CAPA  Rev 4.pdf",
    "Risk Management":                                     f"{_D2}/29 43012 SOP Risk Management Procedure.pdf",
    "Internal Audit Program":                              f"{_D1}/10 43015 SOP, Internal Audit Program.pdf",
    "Customer Feedback, Complaint Handling and Reporting": f"{_D1}/13 43019 SOP Customer Feedback Complaint Handling and Reporting Ver 3.pdf",
    "Purchasing Process":                                  f"{_D3}/47 43032 SOP Purchasing Process_Ver4 .pdf",
    "Vendor Management":                                   f"{_D3}/49 43010 SOP Vendor Managment.pdf",
    "Validation Procedure":                                f"{_D3}/55 43018 SOP, Validation Procedure_Rev 4 JLM (2).pdf",
    "Design and Development":                              f"{_D4}/84 43028 SOP Design and Development Rev 4.pdf",
    "Design Changes":                                      f"{_D4}/85 43033 SOP, Design Changes.pdf",
    "Nonconforming Materials":                             f"{_D3}/39 43009 SOP Nonconforming Materials.pdf",
    "Training":                                            f"{_D1}/17 43017 SOP Training.pdf",
    "Medical Device Reporting":                            f"{_D1}/14 44002 SOP, Medical Device Reporting Rev 1.pdf",
    "Corrections and Removals":                            f"{_D1}/16 44000 SOP Corrections and Removals Rev 2.pdf",
    "Manufacturing Processes and Material Handling":       f"{_D3}/54 41006 SOP Mfg Processes and Material Handling (8).pdf",
    "Equipment Controls":                                  f"{_D3}/58 41005 SOP.pdf",
    "Environmental Controls and Monitoring":               f"{_D3}/56 41008 SOP Imbed Environmental Controls and Monitoring.pdf",
    "Sterilization Process Monitoring and Control":        f"{_D3}/57 43026 SOP Sterilization Process Routine Monitoring and Control .pdf",
    "Line Clearance":                                      f"{_D4}/71 41001 SOP Line Clearance Rev 0.pdf",
    "Gowning in Manufacturing Area":                       f"{_D4}/80 41039 SOP Gowning in the Manufacturing Area.pdf",
    "Return Material Authorization":                       f"{_D3}/38 47031 WI Return Material Authorization (RMA).pdf",
}

QM_SECTIONS = [
    "1 INTRODUCTION",
    "2 SCOPE",
    "3 DEFINITIONS AND ABBREVIATIONS",
    "4 QUALITY MANAGEMENT SYSTEM",
    "5 MANAGEMENT RESPONSIBILITY",
    "6 RESOURCE MANAGEMENT",
    "7 PRODUCT REALIZATION",
    "8 MEASUREMENT, ANALYSIS AND IMPROVEMENT",
]

SOP_SECTIONS = [
    "1 PURPOSE AND SCOPE",
    "2 RESPONSIBILITIES",
    "3 DEFINITIONS",
    "4 REFERENCES",
    "5 FORMS, RECORDS AND ATTACHMENTS",
    "6 PROCEDURE",
]

FDA_TABLE = [
    ["21 CFR Part 820", "Requirement", "Company Implementation"],
    ["820.20", "Management Responsibility", "Quality Policy, Management Review, Org Chart"],
    ["820.22", "Quality Audit", "Annual internal audit program, SOP-AUD-001"],
    ["820.25", "Personnel", "Training program, competency records, SOP-TRN-001"],
    ["820.30", "Design Controls", "Design control procedure, DHF, SOP-DC-001"],
    ["820.40", "Document Controls", "Document control SOP, revision history, SOP-DOC-001"],
    ["820.50", "Purchasing Controls", "Supplier qualification, AVL, SOP-SUP-001"],
    ["820.60", "Identification", "Product labeling and traceability procedure"],
    ["820.65", "Traceability", "Lot traceability for sterile wound dressings"],
    ["820.70", "Production Controls", "Manufacturing SOPs, environmental monitoring"],
    ["820.72", "Inspection/Test Equipment", "Calibration program, equipment log"],
    ["820.75", "Process Validation", "Sterilization, sealing, coating validation"],
    ["820.80", "Acceptance Activities", "Incoming, in-process, finished product inspection"],
    ["820.90", "Nonconforming Product", "NCR procedure, MRB process, SOP-NC-001"],
    ["820.100", "CAPA", "CAPA procedure, effectiveness verification, SOP-CAPA-001"],
    ["820.198", "Complaint Handling", "Customer complaint SOP, MDR reporting, SOP-CMP-001"],
]

ISO_TABLE = [
    ["ISO 13485:2016 Clause", "Title", "Status", "Reference"],
    ["4.1", "General QMS Requirements", "Compliant", "QM-001"],
    ["4.2", "Documentation Requirements", "Compliant", "SOP-DOC-001"],
    ["5.1", "Management Commitment", "Compliant", "Quality Policy"],
    ["5.3", "Quality Policy", "Compliant", "QM-001 Sec. 2"],
    ["5.6", "Management Review", "Compliant", "Management Review SOP"],
    ["6.2", "Human Resources", "Compliant", "SOP-TRN-001"],
    ["6.3", "Infrastructure", "Compliant", "Facility & Equipment Plan"],
    ["7.1", "Product Realization Planning", "Compliant", "SOP-DC-001"],
    ["7.3", "Design and Development", "Compliant", "SOP-DC-001, DHF"],
    ["7.4", "Purchasing", "Compliant", "SOP-SUP-001"],
    ["7.5", "Production Controls", "Compliant", "Manufacturing SOPs"],
    ["7.6", "Control of M&ME", "Compliant", "Calibration Program"],
    ["8.2", "Monitoring & Measurement", "Compliant", "SOP-AUD-001"],
    ["8.3", "Nonconforming Product", "Compliant", "SOP-NC-001"],
    ["8.5", "Improvement / CAPA", "Compliant", "SOP-CAPA-001"],
]

MDR_TABLE = [
    ["EU MDR 2017/745", "Requirement", "Company Approach"],
    ["Article 10", "General Obligations of Manufacturers", "QMS per ISO 13485, technical documentation"],
    ["Article 13", "Obligations of Authorized Representatives", "EU AR appointed pre-commercialization"],
    ["Article 14", "General Obligations of Importers", "Importer agreement in place"],
    ["Annex I", "General Safety & Performance Requirements", "GSPR checklist maintained in DHF"],
    ["Annex II", "Technical Documentation", "Full technical file per Annex II structure"],
    ["Annex III", "Post-Market Surveillance", "PMS plan and PSUR schedule established"],
    ["Annex IX Ch.I", "QMS Assessment (Class IIa)", "Notified body QMS audit required"],
    ["Article 83", "Post-Market Surveillance", "PMS SOP, complaint trending, PMCF plan"],
    ["Article 87", "Reporting of Serious Incidents", "Vigilance reporting SOP, EUDAMED registration"],
]

LIGHT_GRAY = colors.HexColor("#F0F0F0")
MID_GRAY = colors.HexColor("#AAAAAA")
TABLE_HEADER = colors.HexColor("#E8E8E8")
# Kept for backward-compat with table styles
NAVY = colors.HexColor("#1B2A4A")
BLUE = colors.HexColor("#4A9FE0")

DOC_TYPES = {
    # ── Quality Management System (mirrors Imbed 42xxx / 43xxx) ──────────────
    "Quality Manual":                                    "QM",
    "Document and Change Control":                       "SOP-DOC",
    "Corrective and Preventive Action (CAPA)":           "SOP-CAPA",
    "Risk Management":                                   "SOP-RM",
    "Internal Audit Program":                            "SOP-AUD",
    "Customer Feedback, Complaint Handling and Reporting": "SOP-CMP",
    "Purchasing Process":                                "SOP-PUR",
    "Vendor Management":                                 "SOP-VND",
    "Validation Procedure":                              "SOP-VAL",
    "Design and Development":                            "SOP-DD",
    "Design Changes":                                    "SOP-DC",
    "Nonconforming Materials":                           "SOP-NC",
    "Training":                                          "SOP-TRN",
    # ── Regulatory (mirrors Imbed 44xxx) ─────────────────────────────────────
    "Medical Device Reporting":                          "SOP-MDR",
    "Corrections and Removals":                          "SOP-CR",
    # ── Manufacturing / Operations (mirrors Imbed 41xxx) ─────────────────────
    "Manufacturing Processes and Material Handling":     "SOP-MFG",
    "Equipment Controls":                                "SOP-EQP",
    "Environmental Controls and Monitoring":             "SOP-ENV",
    "Sterilization Process Monitoring and Control":      "SOP-STER",
    "Line Clearance":                                    "SOP-LC",
    "Gowning in Manufacturing Area":                     "SOP-GOW",
    # ── Work Instructions (mirrors Imbed 47xxx) ───────────────────────────────
    "Return Material Authorization":                     "WI-RMA",
}


def _build_styles():
    custom = {
        "doc_title": ParagraphStyle("doc_title", fontSize=14, textColor=colors.black,
                                    alignment=TA_CENTER, spaceAfter=4, fontName="Helvetica-Bold"),
        "doc_company": ParagraphStyle("doc_company", fontSize=12, textColor=colors.black,
                                      alignment=TA_CENTER, spaceAfter=16, fontName="Helvetica-Bold"),
        "h1": ParagraphStyle("h1", fontSize=11, textColor=colors.black, spaceBefore=14, spaceAfter=4,
                             fontName="Helvetica-Bold"),
        "h2": ParagraphStyle("h2", fontSize=10, textColor=colors.black, spaceBefore=8, spaceAfter=3,
                             fontName="Helvetica-Bold"),
        "h3": ParagraphStyle("h3", fontSize=10, textColor=colors.black, spaceBefore=6, spaceAfter=2,
                             fontName="Helvetica-BoldOblique"),
        "body": ParagraphStyle("body", fontSize=10, textColor=colors.black, spaceAfter=6,
                               fontName="Helvetica", leading=14),
        "bullet": ParagraphStyle("bullet", fontSize=10, textColor=colors.black, spaceAfter=3,
                                 fontName="Helvetica", leftIndent=20, leading=13, bulletIndent=10),
        "toc": ParagraphStyle("toc", fontSize=10, fontName="Helvetica", spaceAfter=3, leading=14),
        "warning": ParagraphStyle("warning", fontSize=8, textColor=colors.black,
                                  alignment=TA_CENTER, fontName="Helvetica-Bold"),
        # legacy aliases kept so other call-sites don't break
        "cover_title": ParagraphStyle("cover_title", fontSize=14, textColor=colors.black,
                                      alignment=TA_CENTER, spaceAfter=4, fontName="Helvetica-Bold"),
        "cover_sub": ParagraphStyle("cover_sub", fontSize=12, textColor=colors.black,
                                    alignment=TA_CENTER, spaceAfter=8, fontName="Helvetica-Bold"),
        "confidential": ParagraphStyle("confidential", fontSize=8, textColor=colors.black,
                                       alignment=TA_CENTER, fontName="Helvetica"),
    }
    return custom


def _make_canvas_class(doc_number: str, doc_version: str, doc_superseded: str):
    """Return a Canvas subclass that renders the Imbed-style document control
    header and confidentiality footer on every page, including 'Page X of Y'."""

    class _DocCanvas(Canvas):
        def __init__(self, *args, **kwargs):
            Canvas.__init__(self, *args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            total = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self._draw_doc_header(total)
                Canvas.showPage(self)
            Canvas.save(self)

        def _draw_doc_header(self, total_pages: int):
            self.saveState()
            w, h = letter
            page_num = self._pageNumber

            # --- Header block ---
            self.setFont("Helvetica", 8)
            self.setFillColor(colors.black)
            line1 = (
                f"Document Control Number: {doc_number}    "
                f"Current Version: {doc_version}    "
                f"Superseded Version: {doc_superseded}    "
                f"Page: {page_num} of {total_pages}"
            )
            self.drawCentredString(w / 2, h - 0.45 * inch, line1)
            self.setFont("Helvetica-Bold", 7.5)
            self.drawCentredString(
                w / 2, h - 0.60 * inch,
                "DO NOT DISSEMINATE. PRINTED COPIES ARE UNCONTROLLED. FOR REFERENCE USE ONLY."
            )
            # Header separator line
            self.setStrokeColor(colors.black)
            self.setLineWidth(0.5)
            self.line(0.75 * inch, h - 0.70 * inch, w - 0.75 * inch, h - 0.70 * inch)

            # --- Footer ---
            self.setFont("Helvetica", 8)
            self.setFillColor(colors.black)
            self.line(0.75 * inch, 0.65 * inch, w - 0.75 * inch, 0.65 * inch)
            self.drawCentredString(
                w / 2, 0.45 * inch,
                f"Proprietary & Confidential  —  Page {page_num} of {total_pages}  —  Uncontrolled if Printed"
            )
            self.restoreState()

    return _DocCanvas


# Legacy thin wrapper kept so old call-sites don't break (no-op now)
def _header_footer(canvas, doc, company_name, doc_number, doc_title):
    pass


def _generate_content(doc_type: str, company_profile: dict, context_chunks: list[str]) -> str:
    """Single LLM call to generate document content. Uses fast model for speed."""
    from modules.llm_client import generate_fast
    company_name = company_profile.get("name", "Company")
    stage = company_profile.get("stage", "Prototype")
    product = company_profile.get("product_type", "Medical device")
    regulatory = company_profile.get("regulatory", "FDA / ISO 13485")
    notes = company_profile.get("notes", "")

    context_block = ""
    if context_chunks:
        context_block = "\n\nRelevant parent company knowledge (redacted):\n" + \
            "\n---\n".join(c[:600] for c in context_chunks[:3])

    prompt = f"""You are a regulatory document writer for medical device companies.
Write a professional {doc_type} for the following company.
Use clear, numbered sections. Be specific and practical. Write in regulatory language.
Do not use placeholders — write actual content appropriate for this company.
Never reveal or reconstruct proprietary parent company details marked [PLACEHOLDER].
{context_block}

Company: {company_name}
Stage: {stage}
Product: {product}
Regulatory scope: {regulatory}
Notes: {notes}

Generate a complete {doc_type} with all standard sections required for {regulatory}.
Use this exact format for each section:
SECTION: [Section Title]
[Content paragraphs]

SUBSECTION: [Subsection Title]
[Content]

Keep each section concise but complete. Cover all required regulatory elements."""

    return generate_fast(prompt, max_tokens=MAX_OUTPUT_TOKENS)


def _parse_content(raw: str) -> list[tuple[str, str]]:
    """Parse LLM output into (type, text) tuples: section/subsection/body/bullet."""
    blocks = []
    for line in raw.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("SECTION:"):
            blocks.append(("section", line.replace("SECTION:", "").strip()))
        elif line.startswith("SUBSECTION:"):
            blocks.append(("subsection", line.replace("SUBSECTION:", "").strip()))
        elif re.match(r"^[\-\*•]\s+", line):
            blocks.append(("bullet", re.sub(r"^[\-\*•]\s+", "", line)))
        elif re.match(r"^\d+\.\s+", line):
            blocks.append(("bullet", line))
        elif line.startswith("##"):
            blocks.append(("section", line.replace("##", "").strip()))
        elif line.startswith("#"):
            blocks.append(("section", line.replace("#", "").strip()))
        else:
            blocks.append(("body", line))
    return blocks


def _read_reference_doc(path: str, max_chars: int = 8000) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        text = ""
        for page in reader.pages:
            text += (page.extract_text() or "") + "\n"
            if len(text) >= max_chars:
                break
        return text[:max_chars]
    except Exception:
        return ""


def _styled_table(data: list, col_widths: list, has_header: bool = True):
    t = Table(data, colWidths=col_widths)
    style = [
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT_GRAY]),
    ]
    if has_header:
        style += [
            ("BACKGROUND", (0, 0), (-1, 0), TABLE_HEADER),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ]
    t.setStyle(TableStyle(style))
    return t


def _generate_section(section_name: str, company_profile: dict, reference_context: str) -> str:
    company_name = company_profile.get("name", "Company")
    product = company_profile.get("product_type", "Medical device")
    regulatory = company_profile.get("regulatory", "FDA / ISO 13485")
    stage = company_profile.get("stage", "Prototype")
    notes = company_profile.get("notes", "")

    ctx = reference_context[:1200] if reference_context else ""

    prompt = f"""You are writing section "{section_name}" of a Quality Manual for {company_name}.

Company: {company_name}
Product: {product}
Regulatory scope: {regulatory}
Stage: {stage}
Notes: {notes}

Reference knowledge from an experienced medical device company (redacted — do not reconstruct hidden values):
{ctx}

Write a complete, detailed "{section_name}" section for a real FDA/ISO 13485 Quality Manual.
Be specific to {company_name}'s product ({product}).
Write 4-6 paragraphs of substantive regulatory content. Use formal regulatory language.
Include specific requirements, responsibilities, and procedures relevant to {regulatory}.
Do NOT use placeholders like [INSERT] — write real policy content.
Use SUBSECTION: headers for sub-topics within the section.
"""
    try:
        return safe_generate(prompt, max_tokens=1000)
    except Exception as e:
        return f"This section covers {section_name} requirements for {company_name}. Content generation encountered an error: {e}"


def generate_comprehensive_pdf(
    doc_type: str,
    company_profile: dict,
    source_pdf_path: str | None = None,
    progress_callback=None,
) -> tuple[bytes, str]:
    """Generate a comprehensive PDF formatted to match Imbed document control standards."""
    company_name = company_profile.get("name", "Company")
    doc_prefix = DOC_TYPES.get(doc_type, "DOC")
    doc_number = f"{doc_prefix}-001"
    doc_version = "1"
    doc_superseded = "0"
    doc_date = date.today().strftime("%B %d, %Y")
    safe_name = company_name.replace(" ", "_")
    safe_type = doc_type.replace(" ", "_")
    filename = f"{safe_name}_{safe_type}_Comprehensive.pdf"
    out_path = Path(OUTPUTS_DIR) / "advisor_reports" / filename

    ref_path = source_pdf_path or DOC_REFERENCE_PATHS.get(doc_type, IMBED_QM_PATH)
    if progress_callback:
        progress_callback(f"Reading reference: {Path(ref_path).name}...")
    reference_text = _read_reference_doc(ref_path)

    styles = _build_styles()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.85 * inch, rightMargin=0.85 * inch,
        topMargin=1.0 * inch, bottomMargin=0.75 * inch,
        title=f"{company_name} — {doc_type}",
        author=company_name,
    )
    pw = 6.8 * inch
    story = []

    # --- Page 1: Title ---
    story.append(Spacer(1, 0.3 * inch))
    story.append(Paragraph(doc_type.upper(), styles["doc_title"]))
    story.append(Paragraph(f"{company_name.upper()}", styles["doc_company"]))
    story.append(HRFlowable(width="100%", thickness=0.75, color=colors.black))
    story.append(Spacer(1, 0.15 * inch))

    # Document metadata table (matches Imbed style)
    meta_data = [
        ["Document Control Number:", doc_number, "Effective Date:", doc_date],
        ["Current Version:", doc_version, "Product:", company_profile.get("product_type", "Medical Device")],
        ["Superseded Version:", doc_superseded, "Regulatory Scope:", company_profile.get("regulatory", "FDA / ISO 13485")],
        ["Prepared by:", "[NAME], [TITLE]", "Approved by:", "[NAME], [TITLE]"],
    ]
    mt = Table(meta_data, colWidths=[1.8*inch, 1.7*inch, 1.6*inch, 1.7*inch])
    mt.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
        ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
        ("FONTNAME", (3, 0), (3, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [LIGHT_GRAY, colors.white]),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(mt)
    story.append(Spacer(1, 0.2 * inch))

    # Revision history table (matches Imbed style)
    story.append(Paragraph("REVISION HISTORY", styles["h1"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black))
    story.append(Spacer(1, 0.05 * inch))
    rev_data = [
        ["Rev", "Date", "Description of Change", "Author", "Approver"],
        [doc_version, doc_date, "Initial release", "[NAME]", "[NAME]"],
    ]
    story.append(_styled_table(rev_data, [0.5*inch, 1.2*inch, 2.8*inch, 1.15*inch, 1.15*inch]))
    story.append(PageBreak())

    # --- Content sections ---
    sections = QM_SECTIONS if doc_type == "Quality Manual" else SOP_SECTIONS
    total = len(sections)

    for i, section_name in enumerate(sections):
        if progress_callback:
            progress_callback(f"Generating section {i+1}/{total}: {section_name}...")

        story.append(Paragraph(section_name, styles["h1"]))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black))
        story.append(Spacer(1, 0.05 * inch))

        raw = _generate_section(section_name, company_profile, reference_text)
        for kind, text in _parse_content(raw):
            if not text.strip():
                continue
            if kind in ("section", "subsection"):
                story.append(Paragraph(text, styles["h2"]))
            elif kind == "bullet":
                story.append(Paragraph(f"• {text}", styles["bullet"]))
            else:
                story.append(Paragraph(text, styles["body"]))
        story.append(Spacer(1, 0.1 * inch))

    # --- Regulatory tables (Quality Manual only) ---
    if doc_type == "Quality Manual":
        if progress_callback:
            progress_callback("Building regulatory compliance tables...")

        # Build FDA table with company name substituted
        fda_table = [FDA_TABLE[0]] + [
            [r[0], r[1], r[2].replace("Kytosan Bio", company_name)] for r in FDA_TABLE[1:]
        ]
        mdr_table = [MDR_TABLE[0]] + [
            [r[0], r[1], r[2].replace("Kytosan Bio", company_name)] for r in MDR_TABLE[1:]
        ]

        story.append(PageBreak())
        story.append(Paragraph("9 FDA 21 CFR PART 820 COMPLIANCE MATRIX", styles["h1"]))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black))
        story.append(Spacer(1, 0.08 * inch))
        story.append(Paragraph(
            f"The following matrix maps FDA 21 CFR Part 820 requirements to "
            f"{company_name}'s quality management system.", styles["body"]))
        story.append(Spacer(1, 0.08 * inch))
        story.append(_styled_table(fda_table, [1.0*inch, 2.0*inch, 3.8*inch]))

        story.append(PageBreak())
        story.append(Paragraph("10 ISO 13485:2016 CROSS-REFERENCE MATRIX", styles["h1"]))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black))
        story.append(Spacer(1, 0.08 * inch))
        story.append(_styled_table(ISO_TABLE, [1.0*inch, 2.2*inch, 1.2*inch, 2.4*inch]))

        story.append(PageBreak())
        story.append(Paragraph("11 EU MDR 2017/745 REQUIREMENTS", styles["h1"]))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black))
        story.append(Spacer(1, 0.08 * inch))
        story.append(_styled_table(mdr_table, [1.5*inch, 2.4*inch, 2.9*inch]))

        story.append(PageBreak())
        story.append(Paragraph("APPENDIX A — QMS CONTROLLED DOCUMENT LIST", styles["h1"]))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black))
        story.append(Spacer(1, 0.08 * inch))
        proc_data = [["Document Number", "Title", "Type", "Rev"]]
        procedures = [
            ("QM-001", "Quality Manual", "Manual", "1"),
            ("SOP-DOC-001", "Document and Change Control", "SOP", "1"),
            ("SOP-TRN-001", "Training and Competency", "SOP", "1"),
            ("SOP-DC-001", "Design and Development Controls", "SOP", "1"),
            ("SOP-RM-001", "Risk Management", "SOP", "1"),
            ("SOP-SUP-001", "Supplier Qualification and Control", "SOP", "1"),
            ("SOP-CAPA-001", "Corrective and Preventive Action", "SOP", "1"),
            ("SOP-NC-001", "Nonconforming Product Control", "SOP", "1"),
            ("SOP-AUD-001", "Internal Audit Program", "SOP", "1"),
            ("SOP-CMP-001", "Complaint Handling and MDR Reporting", "SOP", "1"),
            ("SOP-VAL-001", "Validation and Verification", "SOP", "1"),
            ("SOP-MFG-001", "Manufacturing Process Controls", "SOP", "1"),
            ("SOP-ENV-001", "Environmental Monitoring", "SOP", "1"),
            ("SOP-STER-001", "Sterilization Process Control", "SOP", "1"),
            ("SOP-CAL-001", "Equipment Calibration", "SOP", "1"),
            ("SOP-LAB-001", "Labeling Controls", "SOP", "1"),
            ("SOP-PMS-001", "Post-Market Surveillance", "SOP", "1"),
            ("SOP-MDR-001", "Medical Device Reporting", "SOP", "1"),
            ("SOP-REC-001", "Record Control", "SOP", "1"),
            ("SOP-MRB-001", "Material Review Board", "SOP", "1"),
        ]
        for row in procedures:
            proc_data.append(list(row))
        story.append(_styled_table(proc_data, [1.4*inch, 3.2*inch, 1.0*inch, 1.2*inch]))

    if progress_callback:
        progress_callback("Formatting and building PDF...")

    canvas_class = _make_canvas_class(doc_number, doc_version, doc_superseded)
    doc.build(story, canvasmaker=canvas_class)
    pdf_bytes = buf.getvalue()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_bytes(pdf_bytes)

    if progress_callback:
        progress_callback(f"Done — saved as {filename}")

    return pdf_bytes, filename


def generate_pdf(
    doc_type: str,
    company_profile: dict,
    context_chunks: list[str] | None = None,
    progress_callback=None,
) -> tuple[bytes, str]:
    """
    Generate a PDF document. Returns (pdf_bytes, filename).
    """
    context_chunks = list(context_chunks or [])

    # Always prepend the specific reference doc for this doc type so the LLM
    # uses the correct SOP/manual rather than whatever ChromaDB happens to return.
    ref_path = DOC_REFERENCE_PATHS.get(doc_type, IMBED_QM_PATH)
    ref_text = _read_reference_doc(ref_path, max_chars=2000)
    if ref_text:
        context_chunks.insert(0, ref_text)

    company_name = company_profile.get("name", "Company")
    doc_prefix = DOC_TYPES.get(doc_type, "DOC")
    doc_number = f"{doc_prefix}-001"
    doc_version = "1"
    doc_superseded = "0"
    doc_date = date.today().strftime("%B %d, %Y")
    filename = f"{company_name.replace(' ', '_')}_{doc_type.replace(' ', '_')}.pdf"
    out_path = Path(OUTPUTS_DIR) / "advisor_reports" / filename

    if progress_callback:
        progress_callback(f"Generating {doc_type} content via AI...")

    raw_content = _generate_content(doc_type, company_profile, context_chunks)

    if progress_callback:
        progress_callback("Formatting PDF...")

    styles = _build_styles()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.85 * inch, rightMargin=0.85 * inch,
        topMargin=1.0 * inch, bottomMargin=0.75 * inch,
        title=f"{company_name} — {doc_type}",
        author=company_name,
    )

    story = []

    # Title block (Imbed style — no decorative cover page)
    story.append(Spacer(1, 0.3 * inch))
    story.append(Paragraph(doc_type.upper(), styles["doc_title"]))
    story.append(Paragraph(company_name.upper(), styles["doc_company"]))
    story.append(HRFlowable(width="100%", thickness=0.75, color=colors.black))
    story.append(Spacer(1, 0.15 * inch))

    meta_data = [
        ["Document Control Number:", doc_number, "Effective Date:", doc_date],
        ["Current Version:", doc_version, "Product:", company_profile.get("product_type", "Medical Device")],
        ["Superseded Version:", doc_superseded, "Regulatory Scope:", company_profile.get("regulatory", "FDA / ISO 13485")],
        ["Prepared by:", "[NAME], [TITLE]", "Approved by:", "[NAME], [TITLE]"],
    ]
    mt = Table(meta_data, colWidths=[1.8*inch, 1.7*inch, 1.6*inch, 1.7*inch])
    mt.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [LIGHT_GRAY, colors.white]),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(mt)
    story.append(Spacer(1, 0.2 * inch))

    # Content sections
    for kind, text in _parse_content(raw_content):
        if not text.strip():
            continue
        if kind == "section":
            story.append(Spacer(1, 0.1 * inch))
            story.append(Paragraph(text.upper(), styles["h1"]))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black))
        elif kind == "subsection":
            story.append(Paragraph(text, styles["h2"]))
        elif kind == "bullet":
            story.append(Paragraph(f"• {text}", styles["bullet"]))
        else:
            story.append(Paragraph(text, styles["body"]))

    canvas_class = _make_canvas_class(doc_number, doc_version, doc_superseded)
    doc.build(story, canvasmaker=canvas_class)
    pdf_bytes = buf.getvalue()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_bytes(pdf_bytes)

    if progress_callback:
        progress_callback(f"Saved: {filename}")

    return pdf_bytes, filename


# ---------------------------------------------------------------------------
# Word (.docx) generator
# ---------------------------------------------------------------------------

def generate_word_doc(
    doc_type: str,
    company_profile: dict,
    comprehensive: bool = False,
    progress_callback=None,
) -> tuple[bytes, str]:
    """Generate a Word (.docx) document. Returns (docx_bytes, filename)."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    company_name = company_profile.get("name", "Company")
    doc_prefix = DOC_TYPES.get(doc_type, "DOC")
    doc_number = f"{doc_prefix}-001"
    rev = "Rev 0.1 DRAFT"
    doc_date = date.today().strftime("%B %d, %Y")
    filename = f"{company_name.replace(' ', '_')}_{doc_type.replace(' ', '_')}.docx"
    out_path = Path(OUTPUTS_DIR) / "advisor_reports" / filename

    d = DocxDocument()

    # Styles helpers
    navy_rgb = RGBColor(0x1B, 0x2A, 0x4A)
    blue_rgb = RGBColor(0x4A, 0x9F, 0xE0)

    def _h1(text):
        p = d.add_heading(text, level=1)
        p.runs[0].font.color.rgb = navy_rgb

    def _h2(text):
        p = d.add_heading(text, level=2)
        p.runs[0].font.color.rgb = blue_rgb

    def _body(text):
        p = d.add_paragraph(text)
        p.style.font.size = Pt(10)

    def _bullet(text):
        d.add_paragraph(text, style="List Bullet")

    # Cover page
    title_p = d.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(company_name)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = navy_rgb

    sub_p = d.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_p.add_run(doc_type)
    run2.font.size = Pt(16)
    run2.font.color.rgb = blue_rgb

    d.add_paragraph()
    meta = [
        ("Document Number", doc_number),
        ("Revision", rev),
        ("Effective Date", doc_date),
        ("Regulatory Scope", company_profile.get("regulatory", "FDA / ISO 13485")),
        ("Product", company_profile.get("product_type", "Medical Device")),
        ("Stage", company_profile.get("stage", "")),
    ]
    tbl = d.add_table(rows=len(meta), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(meta):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = value
        tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    d.add_page_break()

    # Sections
    ref_path = DOC_REFERENCE_PATHS.get(doc_type, IMBED_QM_PATH)
    reference_text = _read_reference_doc(ref_path) if comprehensive else ""
    sections = QM_SECTIONS if doc_type == "Quality Manual" else [
        "1. Purpose and Scope", "2. Responsibilities", "3. Definitions",
        "4. Procedure", "5. Records", "6. References",
    ]
    total = len(sections)

    # For Quick Template, generate all content in one LLM call (same as PDF path)
    if not comprehensive:
        if progress_callback:
            progress_callback("Generating document content...")
        quick_content = _generate_content(doc_type, company_profile, [])

    for i, section_name in enumerate(sections):
        if progress_callback and comprehensive:
            progress_callback(f"Generating section {i+1}/{total}: {section_name}...")
        _h1(section_name)
        raw = _generate_section(section_name, company_profile, reference_text) if comprehensive \
            else quick_content
        for kind, text in _parse_content(raw):
            if not text.strip():
                continue
            if kind in ("section", "subsection"):
                _h2(text)
            elif kind == "bullet":
                _bullet(text)
            else:
                _body(text)

    # Regulatory tables for Quality Manual
    if doc_type == "Quality Manual":
        if progress_callback:
            progress_callback("Adding regulatory compliance tables...")
        d.add_page_break()
        _h1("13. FDA 21 CFR Part 820 Compliance Matrix")
        t = d.add_table(rows=len(FDA_TABLE), cols=3)
        t.style = "Table Grid"
        for r, row in enumerate(FDA_TABLE):
            for c, val in enumerate(row):
                cell = t.rows[r].cells[c]
                cell.text = val
                if r == 0:
                    cell.paragraphs[0].runs[0].bold = True

        d.add_page_break()
        _h1("14. EU MDR 2017/745 Requirements")
        t2 = d.add_table(rows=len(MDR_TABLE), cols=3)
        t2.style = "Table Grid"
        for r, row in enumerate(MDR_TABLE):
            for c, val in enumerate(row):
                cell = t2.rows[r].cells[c]
                cell.text = val
                if r == 0:
                    cell.paragraphs[0].runs[0].bold = True

        d.add_page_break()
        _h1("15. ISO 13485:2016 Cross-Reference")
        t3 = d.add_table(rows=len(ISO_TABLE), cols=4)
        t3.style = "Table Grid"
        for r, row in enumerate(ISO_TABLE):
            for c, val in enumerate(row):
                cell = t3.rows[r].cells[c]
                cell.text = val
                if r == 0:
                    cell.paragraphs[0].runs[0].bold = True

    buf = io.BytesIO()
    d.save(buf)
    docx_bytes = buf.getvalue()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_bytes(docx_bytes)
    if progress_callback:
        progress_callback(f"Saved: {filename}")
    return docx_bytes, filename


# ---------------------------------------------------------------------------
# Excel (.xlsx) generator
# ---------------------------------------------------------------------------

def generate_excel_doc(
    doc_type: str,
    company_profile: dict,
    progress_callback=None,
) -> tuple[bytes, str]:
    """Generate an Excel (.xlsx) workbook. Returns (xlsx_bytes, filename)."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    company_name = company_profile.get("name", "Company")
    doc_prefix = DOC_TYPES.get(doc_type, "DOC")
    doc_number = f"{doc_prefix}-001"
    doc_date = date.today().strftime("%B %d, %Y")
    filename = f"{company_name.replace(' ', '_')}_{doc_type.replace(' ', '_')}.xlsx"
    out_path = Path(OUTPUTS_DIR) / "advisor_reports" / filename

    wb = openpyxl.Workbook()
    navy_hex = "1B2A4A"
    blue_hex = "4A9FE0"
    light_gray = "F5F5F5"
    white = "FFFFFF"

    navy_fill = PatternFill("solid", fgColor=navy_hex)
    blue_fill = PatternFill("solid", fgColor=blue_hex)
    gray_fill = PatternFill("solid", fgColor=light_gray)
    header_font = Font(bold=True, color=white, size=11)
    title_font = Font(bold=True, color=navy_hex, size=13)
    bold_navy = Font(bold=True, color=navy_hex)
    thin = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    wrap = Alignment(wrap_text=True, vertical="top")

    def _header_row(ws, row_data, row_num, fill=None):
        fill = fill or navy_fill
        for col, val in enumerate(row_data, 1):
            c = ws.cell(row=row_num, column=col, value=val)
            c.font = header_font
            c.fill = fill
            c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            c.border = border

    def _data_row(ws, row_data, row_num, alternate=False):
        fill = gray_fill if alternate else PatternFill("solid", fgColor=white)
        for col, val in enumerate(row_data, 1):
            c = ws.cell(row=row_num, column=col, value=val)
            c.fill = fill
            c.alignment = wrap
            c.border = border

    def _set_col_widths(ws, widths):
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    # --- Sheet 1: Document Info ---
    ws1 = wb.active
    ws1.title = "Document Info"
    ws1.merge_cells("A1:D1")
    title_cell = ws1["A1"]
    title_cell.value = f"{company_name} — {doc_type}"
    title_cell.font = Font(bold=True, color=white, size=14)
    title_cell.fill = navy_fill
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws1.row_dimensions[1].height = 30

    info = [
        ("Document Number", doc_number), ("Revision", "Rev 0.1 DRAFT"),
        ("Effective Date", doc_date), ("Company", company_name),
        ("Product", company_profile.get("product_type", "")),
        ("Regulatory Scope", company_profile.get("regulatory", "")),
        ("Stage", company_profile.get("stage", "")),
        ("Prepared by", "[NAME], [TITLE]"), ("Approved by", "[NAME], [TITLE]"),
    ]
    for i, (label, value) in enumerate(info, 2):
        ws1.cell(row=i, column=1, value=label).font = bold_navy
        ws1.cell(row=i, column=2, value=value)
    _set_col_widths(ws1, [22, 40, 20, 20])

    if progress_callback:
        progress_callback("Building FDA compliance sheet...")

    # --- Sheet 2: FDA 21 CFR Part 820 ---
    ws2 = wb.create_sheet("FDA 21 CFR 820")
    ws2.merge_cells("A1:C1")
    ws2["A1"].value = "FDA 21 CFR Part 820 Compliance Matrix"
    ws2["A1"].font = Font(bold=True, color=white, size=12)
    ws2["A1"].fill = navy_fill
    ws2["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws2.row_dimensions[1].height = 25
    _header_row(ws2, FDA_TABLE[0], 2, blue_fill)
    for i, row in enumerate(FDA_TABLE[1:], 3):
        _data_row(ws2, row, i, alternate=(i % 2 == 0))
    _set_col_widths(ws2, [14, 28, 42])

    if progress_callback:
        progress_callback("Building ISO 13485 sheet...")

    # --- Sheet 3: ISO 13485 ---
    ws3 = wb.create_sheet("ISO 13485-2016")
    ws3.merge_cells("A1:D1")
    ws3["A1"].value = "ISO 13485:2016 Cross-Reference Matrix"
    ws3["A1"].font = Font(bold=True, color=white, size=12)
    ws3["A1"].fill = navy_fill
    ws3["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws3.row_dimensions[1].height = 25
    _header_row(ws3, ISO_TABLE[0], 2, blue_fill)
    for i, row in enumerate(ISO_TABLE[1:], 3):
        _data_row(ws3, row, i, alternate=(i % 2 == 0))
    _set_col_widths(ws3, [14, 30, 14, 28])

    if progress_callback:
        progress_callback("Building EU MDR sheet...")

    # --- Sheet 4: EU MDR ---
    ws4 = wb.create_sheet("EU MDR 2017-745")
    ws4.merge_cells("A1:C1")
    ws4["A1"].value = "EU MDR 2017/745 Requirements"
    ws4["A1"].font = Font(bold=True, color=white, size=12)
    ws4["A1"].fill = navy_fill
    ws4["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws4.row_dimensions[1].height = 25
    _header_row(ws4, MDR_TABLE[0], 2, blue_fill)
    for i, row in enumerate(MDR_TABLE[1:], 3):
        _data_row(ws4, row, i, alternate=(i % 2 == 0))
    _set_col_widths(ws4, [18, 34, 36])

    if progress_callback:
        progress_callback("Building QMS document list sheet...")

    # --- Sheet 5: QMS Procedure List ---
    ws5 = wb.create_sheet("QMS Document List")
    ws5.merge_cells("A1:E1")
    ws5["A1"].value = f"{company_name} — QMS Controlled Document List"
    ws5["A1"].font = Font(bold=True, color=white, size=12)
    ws5["A1"].fill = navy_fill
    ws5["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws5.row_dimensions[1].height = 25
    proc_headers = ["Doc Number", "Title", "Type", "Revision", "Status"]
    _header_row(ws5, proc_headers, 2, blue_fill)
    procedures = [
        ("QM-001", "Quality Manual", "Manual", "0.1", "DRAFT"),
        ("SOP-DOC-001", "Document and Change Control", "SOP", "0.1", "DRAFT"),
        ("SOP-TRN-001", "Training and Competency", "SOP", "0.1", "DRAFT"),
        ("SOP-DC-001", "Design and Development Controls", "SOP", "0.1", "DRAFT"),
        ("SOP-RM-001", "Risk Management", "SOP", "0.1", "DRAFT"),
        ("SOP-SUP-001", "Supplier Qualification and Control", "SOP", "0.1", "DRAFT"),
        ("SOP-CAPA-001", "Corrective and Preventive Action", "SOP", "0.1", "DRAFT"),
        ("SOP-NC-001", "Nonconforming Product Control", "SOP", "0.1", "DRAFT"),
        ("SOP-AUD-001", "Internal Audit Program", "SOP", "0.1", "DRAFT"),
        ("SOP-CMP-001", "Complaint Handling and MDR Reporting", "SOP", "0.1", "DRAFT"),
        ("SOP-VAL-001", "Validation and Verification", "SOP", "0.1", "DRAFT"),
        ("SOP-MFG-001", "Manufacturing Process Controls", "SOP", "0.1", "DRAFT"),
        ("SOP-ENV-001", "Environmental Monitoring", "SOP", "0.1", "DRAFT"),
        ("SOP-STER-001", "Sterilization Process Control", "SOP", "0.1", "DRAFT"),
        ("SOP-CAL-001", "Equipment Calibration", "SOP", "0.1", "DRAFT"),
        ("SOP-LAB-001", "Labeling Controls", "SOP", "0.1", "DRAFT"),
        ("SOP-PMS-001", "Post-Market Surveillance", "SOP", "0.1", "DRAFT"),
        ("SOP-MDR-001", "Medical Device Reporting (MDR/Vigilance)", "SOP", "0.1", "DRAFT"),
        ("SOP-REC-001", "Record Control", "SOP", "0.1", "DRAFT"),
        ("SOP-MRB-001", "Material Review Board", "SOP", "0.1", "DRAFT"),
    ]
    for i, row in enumerate(procedures, 3):
        _data_row(ws5, list(row), i, alternate=(i % 2 == 0))
    _set_col_widths(ws5, [16, 40, 10, 10, 10])

    buf = io.BytesIO()
    wb.save(buf)
    xlsx_bytes = buf.getvalue()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_bytes(xlsx_bytes)
    if progress_callback:
        progress_callback(f"Saved: {filename}")
    return xlsx_bytes, filename
