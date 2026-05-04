"""
Generate the Kytosan Bio Quality Manual as a Word (.docx) document.
Run with: /Users/manika/My Drive/VentureStudio/.venv/bin/python3 generate_kytosan_qm.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = Path(
    "/Users/manika/My Drive/VentureStudio/venture_studio_ai/outputs/advisor_reports/Kytosan_Bio_Quality_Manual.docx"
)

# ── colour constants ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x2A, 0x4A)   # heading 1
DARK_BLUE = RGBColor(0x1F, 0x49, 0x7D) # heading 2
BLACK  = RGBColor(0x00, 0x00, 0x00)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xD9, 0xD9, 0xD9)
MID_BLUE   = RGBColor(0x2E, 0x74, 0xB5)  # table header fill


# ── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb_hex: str):
    """Set table cell background colour (e.g. '1B2A4A')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex)
    tcPr.append(shd)


def set_cell_borders(table, border_color="1B2A4A"):
    """Apply thin borders to every cell in a table."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), border_color)
                tcBorders = tcPr.find(qn("w:tcBorders"))
                if tcBorders is None:
                    tcBorders = OxmlElement("w:tcBorders")
                    tcPr.append(tcBorders)
                tcBorders.append(border)


def add_heading1(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    run = p.runs[0]
    run.font.color.rgb = NAVY
    run.font.bold = True
    run.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph(text, style="Heading 2")
    run = p.runs[0]
    run.font.color.rgb = DARK_BLUE
    run.font.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph(text, style="Heading 3")
    run = p.runs[0]
    run.font.color.rgb = MID_BLUE
    run.font.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table_header_row(table, headers, fill_hex="1B2A4A"):
    row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, fill_hex)


def add_table_data_row(table, row_idx, data, shaded=False):
    row = table.rows[row_idx]
    fill = "D9D9D9" if shaded else "FFFFFF"
    for i, val in enumerate(data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = BLACK
        if shaded:
            set_cell_bg(cell, fill)


def add_page_break(doc):
    doc.add_page_break()


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1B2A4A")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ── header / footer ───────────────────────────────────────────────────────────

def add_header_footer(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # ---- Header (non-first pages) ----
    header = section.header
    header.is_linked_to_previous = False
    # Clear default paragraph
    for p in header.paragraphs:
        p.clear()
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Kytosan Bio  |  Quality Manual  |  QM-001  Rev 0.1 DRAFT")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_BLUE
    run.font.italic = True

    # ---- Footer (non-first pages) ----
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = fp.add_run("CONFIDENTIAL — QUALITY MANUAL    |    Kytosan Bio    |    Page ")
    run2.font.name = "Calibri"
    run2.font.size = Pt(9)
    run2.font.color.rgb = DARK_BLUE

    # Add page number field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run3 = fp.add_run()
    run3.font.name = "Calibri"
    run3.font.size = Pt(9)
    run3.font.color.rgb = DARK_BLUE
    run3._r.append(fldChar1)
    run3._r.append(instrText)
    run3._r.append(fldChar2)
    run3._r.append(fldChar3)

    run4 = fp.add_run("  |  Uncontrolled if Printed")
    run4.font.name = "Calibri"
    run4.font.size = Pt(9)
    run4.font.color.rgb = DARK_BLUE


# ═════════════════════════════════════════════════════════════════════════════
# MAIN DOCUMENT BUILDER
# ═════════════════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.5)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    add_header_footer(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # COVER PAGE
    # ─────────────────────────────────────────────────────────────────────────
    doc.add_paragraph()  # spacing
    doc.add_paragraph()

    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_logo.add_run("KYTOSAN BIO")
    r.font.name = "Calibri"
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = NAVY

    doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_title.add_run("QUALITY MANUAL")
    r2.font.name = "Calibri"
    r2.font.size = Pt(32)
    r2.font.bold = True
    r2.font.color.rgb = NAVY

    add_horizontal_rule(doc)

    doc.add_paragraph()

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_sub.add_run("Antimicrobial Wound Dressing — Chitosan Polymer Matrix")
    r3.font.name = "Calibri"
    r3.font.size = Pt(14)
    r3.font.italic = True
    r3.font.color.rgb = DARK_BLUE

    doc.add_paragraph()
    doc.add_paragraph()

    # Cover metadata table
    cov_tbl = doc.add_table(rows=6, cols=2)
    cov_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cov_tbl.style = "Table Grid"
    meta_rows = [
        ("Document Number",   "QM-001"),
        ("Revision",          "Rev 0.1 DRAFT"),
        ("Effective Date",    "May 4, 2026"),
        ("Prepared By",       "[NAME], [TITLE]"),
        ("Reviewed By",       "[NAME], [TITLE]"),
        ("Approved By",       "[NAME], CEO / Management Representative"),
    ]
    for i, (label, val) in enumerate(meta_rows):
        row = cov_tbl.rows[i]
        row.cells[0].text = label
        row.cells[1].text = val
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(11)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_borders(cov_tbl)

    doc.add_paragraph()
    doc.add_paragraph()

    p_conf = doc.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = p_conf.add_run(
        "CONFIDENTIAL — This document contains proprietary information of Kytosan Bio.\n"
        "Unauthorized disclosure, reproduction, or distribution is strictly prohibited.\n"
        "Uncontrolled if Printed."
    )
    rc.font.name = "Calibri"
    rc.font.size = Pt(10)
    rc.font.italic = True
    rc.font.color.rgb = DARK_BLUE

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # TABLE OF CONTENTS (manual)
    # ─────────────────────────────────────────────────────────────────────────
    add_heading1(doc, "Table of Contents")

    toc_entries = [
        ("1",    "Introduction and Company Overview", ""),
        ("2",    "Quality Policy", ""),
        ("3",    "Scope of the Quality Management System", ""),
        ("4",    "Normative References", ""),
        ("5",    "Terms and Definitions", ""),
        ("6",    "Quality Management System Requirements", ""),
        ("6.1",  "General Requirements", ""),
        ("6.2",  "Documentation Requirements", ""),
        ("6.3",  "Document Control", ""),
        ("6.4",  "Record Control", ""),
        ("7",    "Management Responsibility", ""),
        ("7.1",  "Management Commitment", ""),
        ("7.2",  "Customer Focus", ""),
        ("7.3",  "Quality Policy (Detail)", ""),
        ("7.4",  "Planning", ""),
        ("7.5",  "Responsibility, Authority, and Communication", ""),
        ("7.6",  "Management Review", ""),
        ("8",    "Resource Management", ""),
        ("8.1",  "Provision of Resources", ""),
        ("8.2",  "Human Resources (Training and Competency)", ""),
        ("8.3",  "Infrastructure", ""),
        ("8.4",  "Work Environment and Contamination Control", ""),
        ("9",    "Product Realization", ""),
        ("9.1",  "Planning of Product Realization", ""),
        ("9.2",  "Customer-Related Processes", ""),
        ("9.3",  "Design and Development Controls", ""),
        ("9.4",  "Purchasing and Supplier Controls", ""),
        ("9.5",  "Production and Service Controls", ""),
        ("9.6",  "Control of Monitoring and Measuring Equipment", ""),
        ("10",   "Measurement, Analysis, and Improvement", ""),
        ("10.1", "General", ""),
        ("10.2", "Monitoring and Measurement", ""),
        ("10.3", "Control of Nonconforming Product", ""),
        ("10.4", "Analysis of Data", ""),
        ("10.5", "Improvement (CAPA and Preventive Action)", ""),
        ("11",   "Regulatory Compliance", ""),
        ("11.1", "FDA 21 CFR Part 820 Compliance Summary", ""),
        ("11.2", "EU MDR 2017/745 Requirements (Class IIa, CE Mark)", ""),
        ("11.3", "ISO 13485:2016 Cross-Reference Matrix", ""),
        ("12",   "Document Control List", ""),
        ("",     "Appendix A — Quality Policy Statement", ""),
        ("",     "Appendix B — Organizational Structure", ""),
        ("",     "Appendix C — Process Map / Turtle Diagram", ""),
    ]

    toc_tbl = doc.add_table(rows=len(toc_entries), cols=2)
    toc_tbl.style = "Table Grid"
    for i, (num, title, _pg) in enumerate(toc_entries):
        c0 = toc_tbl.rows[i].cells[0]
        c1 = toc_tbl.rows[i].cells[1]
        c0.text = num
        c1.text = title
        for cell in (c0, c1):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)
                    if num and "." not in num and num.isdigit():
                        run.font.bold = True
                    if not num:
                        run.font.italic = True
    set_cell_borders(toc_tbl)

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 1 — INTRODUCTION AND COMPANY OVERVIEW
    # ─────────────────────────────────────────────────────────────────────────
    add_heading1(doc, "1. Introduction and Company Overview")

    add_body(doc,
        "Kytosan Bio is an early-stage medical device company dedicated to the design, development, "
        "and manufacturing of innovative wound care solutions. The company's flagship product is an "
        "antimicrobial wound dressing based on a polymer matrix incorporating Chitosan — a naturally "
        "derived biopolymer obtained from chitin (found in crustacean shells) — that exhibits "
        "inherent antimicrobial properties, biocompatibility, and wound-healing promotion capabilities.")

    add_body(doc,
        "This Quality Manual serves as the executive-level summary of the Kytosan Bio Quality "
        "Management System (QMS). It describes the scope, structure, processes, and governance of "
        "the QMS, and demonstrates Kytosan Bio's capability to safely and effectively design, "
        "develop, and manufacture its Chitosan-based antimicrobial wound dressing for burns and "
        "wound management.")

    add_heading2(doc, "1.1 Mission")
    add_body(doc,
        "Kytosan Bio's mission is to improve patient outcomes in wound care by delivering safe, "
        "effective, and clinically validated antimicrobial dressings that harness the natural "
        "properties of Chitosan. The company is committed to maintaining the highest standards of "
        "product quality, regulatory compliance, and continuous improvement across all operations.")

    add_heading2(doc, "1.2 Product Description")
    add_body(doc,
        "The Kytosan Bio wound dressing consists of a biocompatible polymer matrix in which "
        "Chitosan is uniformly distributed. The product is intended for use on partial- and "
        "full-thickness wounds, including burns, surgical wounds, chronic wounds, and traumatic "
        "lacerations. Key product attributes include:")
    add_bullet(doc, "Antimicrobial activity against a broad spectrum of bacterial and fungal pathogens, including MRSA.")
    add_bullet(doc, "Biocompatibility in accordance with ISO 10993 series requirements.")
    add_bullet(doc, "Moist wound healing environment promotion.")
    add_bullet(doc, "Sterile, terminally sterilized presentation (sterilization method to be validated per applicable standards).")
    add_bullet(doc, "Single-use, disposable format.")
    add_bullet(doc, "Available in multiple sizes and configurations to accommodate diverse wound geometries.")

    add_heading2(doc, "1.3 Regulatory Classification and Pathway")
    add_body(doc, "Kytosan Bio's wound dressing is classified and regulated as follows:")

    reg_tbl = doc.add_table(rows=4, cols=3)
    reg_tbl.style = "Table Grid"
    add_table_header_row(reg_tbl, ["Jurisdiction", "Classification", "Regulatory Pathway"])
    rows_data = [
        ("United States (FDA)",       "Class II Medical Device",     "510(k) Premarket Notification"),
        ("European Union (EU MDR)",   "Class IIa Medical Device",    "CE Mark under EU MDR 2017/745, Rule 7"),
        ("Other Markets",             "Subject to local classification", "As applicable per local requirements"),
    ]
    for i, row_d in enumerate(rows_data):
        add_table_data_row(reg_tbl, i + 1, row_d, shaded=(i % 2 == 1))
    set_cell_borders(reg_tbl)
    doc.add_paragraph()

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 2 — QUALITY POLICY
    # ─────────────────────────────────────────────────────────────────────────
    add_heading1(doc, "2. Quality Policy")

    add_body(doc,
        "The Kytosan Bio Management Team has established, approved, and is committed to the "
        "following Quality Policy, which is communicated, understood, and maintained at all "
        "levels of the organization:", bold=True)

    p_policy = doc.add_paragraph()
    p_policy.paragraph_format.left_indent = Inches(0.5)
    p_policy.paragraph_format.right_indent = Inches(0.5)
    p_policy.paragraph_format.space_before = Pt(10)
    p_policy.paragraph_format.space_after = Pt(10)
    rp = p_policy.add_run(
        "\"Kytosan Bio is dedicated to improving patient outcomes through the development and "
        "supply of safe, effective, and innovative Chitosan-based wound care solutions. We are "
        "committed to continuously improving our products and processes by maintaining rigorous "
        "standards of product quality from raw material receipt through finished product release; "
        "ensuring full compliance with applicable regulatory requirements including FDA 21 CFR "
        "Part 820, ISO 13485:2016, and EU MDR 2017/745; and sustaining an effective Quality "
        "Management System that is regularly reviewed for suitability, adequacy, and compliance.\""
    )
    rp.font.name = "Calibri"
    rp.font.size = Pt(11)
    rp.font.italic = True
    rp.font.color.rgb = NAVY

    add_heading2(doc, "2.1 Quality Objectives")
    add_body(doc,
        "The following measurable Quality Objectives are established consistent with the Quality Policy "
        "and are reviewed at each Management Review:")

    add_body(doc, "Patient Safety and Product Efficacy:", bold=True)
    add_bullet(doc, "Complete 100% of final release acceptance activities for each production lot within three business days of production completion.")
    add_bullet(doc, "Achieve zero critical nonconformances related to sterility or biocompatibility in finished product.")
    add_bullet(doc, "Review and evaluate 100% of customer complaints within five business days of receipt.")

    add_body(doc, "Product Quality:", bold=True)
    add_bullet(doc, "Maintain 100% on-time equipment calibration and preventive maintenance for all manufacturing and measurement equipment.")
    add_bullet(doc, "Ensure current qualification status for 100% of suppliers providing materials or components used in manufacturing.")
    add_bullet(doc, "Complete all sterilization validation and revalidation activities prior to initial distribution and following process changes.")

    add_body(doc, "Regulatory Compliance:", bold=True)
    add_bullet(doc, "Ensure 100% on-time training for all Kytosan Bio personnel, aligned with roles and responsibilities, prior to performing work affecting product quality.")
    add_bullet(doc, "Maintain the CAPA system to drive corrective actions and continuous improvement, with no overdue CAPA actions exceeding defined timelines.")
    add_bullet(doc, "Complete all required regulatory submissions and notifications within applicable regulatory timelines.")

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 3 — SCOPE
    # ─────────────────────────────────────────────────────────────────────────
    add_heading1(doc, "3. Scope of the Quality Management System")

    add_body(doc,
        "Kytosan Bio operates and maintains a documented Quality Management System (QMS) designed "
        "to comply with all applicable regulatory requirements and industry standards for the design, "
        "development, and manufacturing of its Chitosan-based antimicrobial wound dressing.")

    add_heading2(doc, "3.1 Product and Process Scope")
    add_body(doc,
        "The QMS encompasses all activities related to the lifecycle of the Kytosan Bio antimicrobial "
        "wound dressing, including:")
    add_bullet(doc, "Research and design and development activities.")
    add_bullet(doc, "Raw material and component procurement and incoming inspection.")
    add_bullet(doc, "Manufacturing, assembly, and in-process controls.")
    add_bullet(doc, "Sterilization and sterile barrier system processes.")
    add_bullet(doc, "Quality control and final product release.")
    add_bullet(doc, "Labeling and packaging.")
    add_bullet(doc, "Storage and distribution.")
    add_bullet(doc, "Post-market surveillance and customer feedback.")
    add_bullet(doc, "Complaint handling and adverse event reporting.")

    add_heading2(doc, "3.2 Applicable Standards and Regulations")
    add_body(doc, "The Kytosan Bio QMS is designed to comply with, at minimum:")
    add_bullet(doc, "FDA 21 CFR Part 820 — Quality System Regulation (Quality Management System Requirements) as updated.")
    add_bullet(doc, "FDA 21 CFR Part 801 — Labeling.")
    add_bullet(doc, "FDA 21 CFR Part 803 — Medical Device Reporting.")
    add_bullet(doc, "FDA 21 CFR Part 806 — Medical Device Corrections and Removals.")
    add_bullet(doc, "ISO 13485:2016 — Medical Devices: Quality Management Systems — Requirements for Regulatory Purposes.")
    add_bullet(doc, "ISO 14971:2019 — Medical Devices: Application of Risk Management to Medical Devices.")
    add_bullet(doc, "ISO 10993 series — Biological Evaluation of Medical Devices.")
    add_bullet(doc, "EU MDR 2017/745 — EU Medical Device Regulation.")
    add_bullet(doc, "ISO 11135 / ISO 11137 — Sterilization of Healthcare Products (as applicable to selected sterilization method).")
    add_bullet(doc, "ISO 11607 — Packaging for Terminally Sterilized Medical Devices.")
    add_bullet(doc, "IEC 62304 — Medical Device Software Life Cycle Processes (if applicable to any software components).")

    add_heading2(doc, "3.3 Exclusions")
    add_body(doc,
        "The following ISO 13485:2016 clauses are excluded from the Kytosan Bio QMS because they "
        "are not applicable to the company's current product and operations:")

    excl_tbl = doc.add_table(rows=4, cols=2)
    excl_tbl.style = "Table Grid"
    add_table_header_row(excl_tbl, ["ISO 13485:2016 Clause", "Justification for Exclusion"])
    exclusions = [
        ("7.5.3 — Installation Activities",
         "Kytosan Bio does not design, manufacture, or support products requiring installation at a customer site."),
        ("7.5.4 — Servicing Activities",
         "Kytosan Bio's wound dressings are single-use, disposable devices not subject to servicing."),
        ("7.5.9.2 — Active Implantable / Implantable Devices",
         "Kytosan Bio does not design or manufacture active implantable or implantable medical devices."),
    ]
    for i, (clause, just) in enumerate(exclusions):
        add_table_data_row(excl_tbl, i + 1, (clause, just), shaded=(i % 2 == 1))
    set_cell_borders(excl_tbl)
    doc.add_paragraph()

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 4 — NORMATIVE REFERENCES
    # ─────────────────────────────────────────────────────────────────────────
    add_heading1(doc, "4. Normative References")
    add_body(doc,
        "For dated references, the edition cited applies. For undated references, the latest "
        "edition of the referenced document, including any amendments, applies.")

    norm_refs = [
        ("ISO 13485:2016",        "Medical devices — Quality management systems — Requirements for regulatory purposes"),
        ("ISO 14971:2019",        "Medical devices — Application of risk management to medical devices"),
        ("ISO 10993-1:2018",      "Biological evaluation of medical devices — Part 1: Evaluation and testing within a risk management process"),
        ("ISO 10993-5",           "Tests for in vitro cytotoxicity"),
        ("ISO 10993-10",          "Tests for skin sensitization"),
        ("ISO 10993-11",          "Tests for systemic toxicity"),
        ("ISO 11135:2014",        "Sterilization of health-care products — Ethylene oxide (if applicable)"),
        ("ISO 11137",             "Sterilization of health-care products — Radiation (if applicable)"),
        ("ISO 11607-1:2019",      "Packaging for terminally sterilized medical devices — Part 1: Requirements for materials, sterile barrier systems, and packaging systems"),
        ("ISO 11607-2:2019",      "Packaging for terminally sterilized medical devices — Part 2: Validation requirements for forming, sealing, and assembly processes"),
        ("IEC 62304:2006 + A1:2015", "Medical device software — Software life cycle processes (if applicable)"),
        ("EU MDR 2017/745",       "Regulation (EU) 2017/745 of the European Parliament and of the Council on medical devices"),
        ("FDA 21 CFR Part 820",   "Quality System Regulation (Current Good Manufacturing Practice)"),
        ("FDA 21 CFR Part 801",   "Labeling"),
        ("FDA 21 CFR Part 803",   "Medical Device Reporting"),
        ("FDA 21 CFR Part 806",   "Medical Devices — Reports of Corrections and Removals"),
        ("ASTM F2132",            "Standard Specification for Accelerated Aging of Sterile Barriers (if applicable)"),
    ]

    norm_tbl = doc.add_table(rows=len(norm_refs) + 1, cols=2)
    norm_tbl.style = "Table Grid"
    add_table_header_row(norm_tbl, ["Standard / Regulation", "Title"])
    for i, (std, title) in enumerate(norm_refs):
        add_table_data_row(norm_tbl, i + 1, (std, title), shaded=(i % 2 == 0))
    set_cell_borders(norm_tbl)
    doc.add_paragraph()

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 5 — TERMS AND DEFINITIONS
    # ─────────────────────────────────────────────────────────────────────────
    add_heading1(doc, "5. Terms and Definitions")
    add_body(doc,
        "For the purposes of this Quality Manual and the Kytosan Bio QMS, the following terms and "
        "definitions apply, in addition to those defined in ISO 13485:2016 and FDA 21 CFR Part 820.3:")

    terms = [
        ("Biocompatibility",
         "The ability of a material or device to perform its intended function with an appropriate host response in a specific application (ISO 10993-1)."),
        ("Chitosan",
         "A linear polysaccharide derived from the partial deacetylation of chitin (found in crustacean exoskeletons and fungal cell walls), exhibiting antimicrobial, hemostatic, and wound-healing properties."),
        ("Corrective Action",
         "Action taken to eliminate the cause of a detected nonconformity or other undesirable situation."),
        ("Design History File (DHF)",
         "A compilation of records which describes the design history of a finished device (21 CFR 820.3(e))."),
        ("Device History Record (DHR)",
         "A compilation of records containing the production history of a finished device (21 CFR 820.3(f))."),
        ("Device Master Record (DMR)",
         "A compilation of records containing the procedures and specifications for a finished device (21 CFR 820.3(j))."),
        ("Management Representative",
         "An individual appointed by top management who is responsible for ensuring QMS requirements are established, maintained, and reported upon."),
        ("Management Team (MT) / Top Management",
         "The senior leadership group of Kytosan Bio with executive responsibility for the QMS, including the CEO, Head of Quality and Regulatory Affairs, Head of Operations, and Head of Research and Development."),
        ("Nonconforming Product",
         "Product that does not conform to specified requirements."),
        ("Post-Market Surveillance (PMS)",
         "A systematic process for collecting and analyzing data and information from the post-production phase of a device."),
        ("Preventive Action",
         "Action taken to eliminate the cause of a potential nonconformity or other undesirable potential situation."),
        ("Quality Management System (QMS)",
         "The organizational structure, responsibilities, procedures, processes, and resources needed to implement quality management."),
        ("Risk Management",
         "A systematic application of management policies, procedures, and practices to the tasks of analyzing, evaluating, controlling, and monitoring risk (ISO 14971)."),
        ("Sterile Barrier System",
         "The minimum package that prevents ingress of microorganisms and allows aseptic presentation of the product at point of use (ISO 11607-1)."),
        ("510(k)",
         "A premarket submission made to FDA to demonstrate that the device to be marketed is at least as safe and effective as, i.e., substantially equivalent to, a legally marketed predicate device."),
        ("Technical Documentation",
         "Documentation required under EU MDR 2017/745 demonstrating conformity of a device with applicable requirements, including design, safety, performance, and manufacturing information."),
    ]

    for term, defn in terms:
        add_body(doc, term + ":", bold=True)
        add_body(doc, defn)
        doc.add_paragraph()

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 6 — QMS REQUIREMENTS
    # ─────────────────────────────────────────────────────────────────────────
    add_heading1(doc, "6. Quality Management System Requirements")

    add_heading2(doc, "6.1 General Requirements")
    add_body(doc,
        "Kytosan Bio has established, documented, implemented, and maintains an effective Quality "
        "Management System in accordance with the requirements of ISO 13485:2016 and FDA 21 CFR "
        "Part 820. The QMS is continually improved through the use of the Quality Policy, quality "
        "objectives, audit results, data analysis, corrective and preventive actions, and management "
        "review.")
    add_body(doc, "Kytosan Bio ensures:")
    add_bullet(doc, "The processes and roles needed for the QMS are identified, and a risk-based approach is applied.")
    add_bullet(doc, "The sequence and interaction of QMS processes are determined and documented.")
    add_bullet(doc, "Criteria and methods for effective operation and control of QMS processes are defined.")
    add_bullet(doc, "Resources and information necessary to support operation and monitoring of QMS processes are provided.")
    add_bullet(doc, "QMS processes are monitored, measured, and analyzed.")
    add_bullet(doc, "Records are established and maintained to demonstrate required conformance.")
    add_bullet(doc, "Changes to QMS processes are evaluated for their impact and controlled in accordance with regulatory requirements.")
    add_body(doc,
        "Kytosan Bio retains responsibility for compliance with applicable standards and regulations "
        "for any outsourced processes that affect product conformity. Outsourced processes are "
        "monitored and controlled proportionate to the risk involved, and requirements are documented "
        "in written quality agreements.")
    add_body(doc,
        "Computer software applications used in the QMS (including the electronic QMS platform, "
        "if applicable) are validated in accordance with the software validation procedure prior to "
        "initial use and following any significant changes, proportionate to risk.")

    add_heading2(doc, "6.2 Documentation Requirements")
    add_heading3(doc, "6.2.1 General")
    add_body(doc, "The Kytosan Bio QMS documentation hierarchy includes:")
    add_bullet(doc, "Level 1: Quality Manual (QM-001) — This document.")
    add_bullet(doc, "Level 2: Standard Operating Procedures (SOPs) — Detailed procedural instructions.")
    add_bullet(doc, "Level 3: Work Instructions (WIs), Forms, and Specifications — Step-level task guidance.")
    add_bullet(doc, "Level 4: Records — Evidence of QMS execution and product conformance.")

    add_heading3(doc, "6.2.2 Quality Manual")
    add_body(doc,
        "This Quality Manual is the executive summary of the Kytosan Bio QMS. It describes the scope "
        "of the QMS, documents the Quality Policy, references all applicable procedures, and describes "
        "interactions between QMS processes.")

    add_heading3(doc, "6.2.3 Device Master Record (DMR)")
    add_body(doc,
        "A Device Master Record is established and maintained for the Kytosan Bio antimicrobial wound "
        "dressing. The DMR contains or references the complete manufacturing process specifications, "
        "quality requirements, and technical documentation for the product, including:")
    add_bullet(doc, "Device description and intended use.")
    add_bullet(doc, "Product specifications (physical, chemical, biological, and performance).")
    add_bullet(doc, "Manufacturing procedures and work instructions.")
    add_bullet(doc, "Packaging, labeling, and shipping specifications.")
    add_bullet(doc, "Sterilization specifications and process parameters.")
    add_bullet(doc, "Inspection, measurement, and test procedures.")

    add_heading3(doc, "6.2.4 Device History Record (DHR)")
    add_body(doc,
        "Device History Records are established and maintained for each production lot to provide "
        "complete traceability and evidence that the product was manufactured in accordance with the DMR. "
        "DHRs include the quantity manufactured, quantity approved for distribution, lot numbers, and "
        "acceptance activity records.")

    add_heading2(doc, "6.3 Document Control")
    add_body(doc,
        "Procedures are established and maintained to control all QMS documents, including externally "
        "generated documents (standards, regulatory guidance). Controls ensure that:")
    add_bullet(doc, "Documents are reviewed and approved for adequacy by authorized personnel prior to issue.")
    add_bullet(doc, "Current revisions of documents are identified and available at points of use.")
    add_bullet(doc, "Invalid or obsolete documents are removed from use and identified to prevent unintended use.")
    add_bullet(doc, "Changes to controlled documents are reviewed, approved, and communicated.")
    add_bullet(doc, "Records of document changes are maintained.")
    add_body(doc,
        "Document retention times are established and are not less than two years, the lifetime of "
        "the device, or the retention time of associated product quality records, whichever is longer.")

    add_heading2(doc, "6.4 Record Control")
    add_body(doc,
        "Procedures are established and maintained for the identification, storage, protection, "
        "retrieval, retention, and disposition of quality records. Quality records are:")
    add_bullet(doc, "Legible, readily identifiable, and retrievable.")
    add_bullet(doc, "Stored in a suitable environment to prevent damage, deterioration, or loss.")
    add_bullet(doc, "Protected against unauthorized access, alteration, or destruction.")
    add_body(doc,
        "Quality record retention times are established and are not less than two years or the "
        "lifetime of the device, whichever is longer. Confidential patient health information is "
        "protected in accordance with applicable regulatory requirements.")

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 7 — MANAGEMENT RESPONSIBILITY
    # ─────────────────────────────────────────────────────────────────────────
    add_heading1(doc, "7. Management Responsibility")

    add_heading2(doc, "7.1 Management Commitment")
    add_body(doc,
        "The Kytosan Bio Management Team is committed to the development, implementation, and "
        "maintenance of the QMS and to its continual improvement. Evidence of this commitment includes:")
    add_bullet(doc, "Communicating to all employees the importance of meeting customer, statutory, and regulatory requirements.")
    add_bullet(doc, "Establishing and communicating the Quality Policy.")
    add_bullet(doc, "Ensuring measurable quality objectives are established and reviewed.")
    add_bullet(doc, "Conducting periodic Management Reviews of the QMS.")
    add_bullet(doc, "Ensuring adequate resources are provided to implement and maintain the QMS.")
    add_bullet(doc, "Promoting a quality culture and continuous improvement mindset throughout the organization.")

    add_heading2(doc, "7.2 Customer Focus")
    add_body(doc,
        "The Kytosan Bio Management Team ensures that customer requirements — including the needs "
        "of patients, healthcare providers, and regulatory bodies — are determined and met. Procedures "
        "are established for determining customer requirements, reviewing requirements prior to supply "
        "commitment, communicating product information, and capturing customer feedback.")
    add_body(doc,
        "Customer satisfaction is monitored through post-market surveillance activities, complaint "
        "handling, and periodic review of feedback data. Information gathered serves as input into "
        "the risk management process and product realization improvement activities.")

    add_heading2(doc, "7.3 Quality Policy")
    add_body(doc,
        "The Kytosan Bio Quality Policy (see Section 2 and Appendix A) is:")
    add_bullet(doc, "Appropriate to the purpose and context of the organization.")
    add_bullet(doc, "Inclusive of a commitment to comply with applicable requirements and to continually improve QMS effectiveness.")
    add_bullet(doc, "A framework for establishing and reviewing quality objectives.")
    add_bullet(doc, "Communicated and understood throughout the organization.")
    add_bullet(doc, "Reviewed for continuing suitability at each Management Review.")

    add_heading2(doc, "7.4 Planning")
    add_heading3(doc, "7.4.1 Quality Objectives")
    add_body(doc,
        "Quality objectives are established at relevant functions and levels within Kytosan Bio. "
        "Objectives are measurable, consistent with the Quality Policy, and reviewed for "
        "achievement progress at each Management Review. See Section 2.1 for current Quality Objectives.")

    add_heading3(doc, "7.4.2 Quality Management System Planning")
    add_body(doc,
        "The Management Team ensures quality planning is performed to define and document how "
        "quality requirements and objectives are met. This includes the Quality Manual, SOPs, "
        "product realization plans, supplier evaluation processes, and product documentation. "
        "The integrity of the QMS is maintained when changes to the QMS are planned and implemented.")

    add_heading2(doc, "7.5 Responsibility, Authority, and Communication")
    add_heading3(doc, "7.5.1 Responsibility and Authority")
    add_body(doc,
        "The Management Team defines, documents, and communicates employee responsibility and "
        "authority through job descriptions, SOPs, and organizational communications.")

    resp_tbl = doc.add_table(rows=6, cols=2)
    resp_tbl.style = "Table Grid"
    add_table_header_row(resp_tbl, ["Functional Area", "Primary Responsibilities"])
    resp_data = [
        ("Chief Executive Officer (CEO)",
         "Overall strategic direction; QMS leadership; regulatory engagement; resource allocation."),
        ("Head of Quality and Regulatory Affairs (Management Representative)",
         "QMS development and management; 510(k) and CE Mark submissions; internal audits; document control; CAPA; complaint administration; regulatory reporting; post-market surveillance; supplier quality."),
        ("Head of Research and Development",
         "Product design and development; design controls; design verification and validation; risk management; biocompatibility and sterilization program; design transfer."),
        ("Head of Operations / Manufacturing",
         "Process development; production and scheduling; incoming inspection; equipment maintenance and calibration; material control; packaging and shipping; sterilization operations."),
        ("All Personnel",
         "Compliance with applicable QMS procedures; reporting quality issues; participating in training; authority to identify and hold nonconforming product."),
    ]
    for i, (role, resp) in enumerate(resp_data):
        add_table_data_row(resp_tbl, i + 1, (role, resp), shaded=(i % 2 == 0))
    set_cell_borders(resp_tbl)
    doc.add_paragraph()

    add_heading3(doc, "7.5.2 Management Representative")
    add_body(doc,
        "The Management Team has appointed the Head of Quality and Regulatory Affairs as the "
        "Kytosan Bio Management Representative. The Management Representative is responsible for:")
    add_bullet(doc, "Ensuring QMS requirements are effectively established and maintained.")
    add_bullet(doc, "Reporting to top management on QMS performance.")
    add_bullet(doc, "Promoting awareness of regulatory requirements throughout the organization.")
    add_bullet(doc, "Placing products or processes on hold as required.")
    add_bullet(doc, "Representing Kytosan Bio during third-party audits and regulatory inspections.")
    add_bullet(doc, "Ensuring the Quality Policy is approved by top management and communicated at all levels.")

    add_heading3(doc, "7.5.3 Internal Communication")
    add_body(doc,
        "The Management Team ensures communication processes are established and that communication "
        "regarding QMS effectiveness occurs regularly. Communication mechanisms include, but are not "
        "limited to: team meetings, quality metrics dashboards, the CAPA process, management reviews, "
        "and internal training events.")

    add_heading2(doc, "7.6 Management Review")
    add_heading3(doc, "7.6.1 General")
    add_body(doc,
        "The Kytosan Bio Management Team reviews the suitability, adequacy, and effectiveness of "
        "the QMS at a minimum of once per calendar year. Additional ad hoc reviews may be convened "
        "in response to significant quality events, regulatory changes, or business decisions "
        "affecting the QMS. Records of management reviews are maintained.")

    add_heading3(doc, "7.6.2 Review Inputs")
    add_body(doc, "Management review agenda items include, but are not limited to:")
    add_bullet(doc, "Internal audit results and findings.")
    add_bullet(doc, "Process and product quality trends and metrics.")
    add_bullet(doc, "Customer feedback, complaints, and post-market surveillance data.")
    add_bullet(doc, "Nonconformance trends and CAPA status.")
    add_bullet(doc, "Supplier performance data.")
    add_bullet(doc, "Progress on quality objectives from the prior review period.")
    add_bullet(doc, "Regulatory changes or new requirements affecting the QMS.")
    add_bullet(doc, "Status of design and development projects.")
    add_bullet(doc, "Recommendations for QMS improvement.")

    add_heading3(doc, "7.6.3 Review Outputs")
    add_body(doc,
        "Management review minutes are generated to document decisions and actions resulting from "
        "the review, including:")
    add_bullet(doc, "QMS and process improvement actions and owners.")
    add_bullet(doc, "Product improvements to satisfy customer requirements.")
    add_bullet(doc, "Resource requirements identified.")
    add_bullet(doc, "Necessary changes in response to regulatory requirements.")
    add_bullet(doc, "Revised quality objectives for the next review period.")

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 8 — RESOURCE MANAGEMENT
    # ─────────────────────────────────────────────────────────────────────────
    add_heading1(doc, "8. Resource Management")

    add_heading2(doc, "8.1 Provision of Resources")
    add_body(doc,
        "Kytosan Bio determines and provides the resources needed to implement and maintain the QMS, "
        "maintain its effectiveness, enhance customer satisfaction, and meet regulatory requirements. "
        "Resource needs are assessed during quality planning, management reviews, and as part of the "
        "CAPA process.")

    add_heading2(doc, "8.2 Human Resources — Training and Competency")
    add_heading3(doc, "8.2.1 General")
    add_body(doc,
        "All personnel performing work affecting product quality are determined to be competent on "
        "the basis of appropriate education, skills, background, training, and experience, as defined "
        "by job description requirements. Consultants providing advice regarding design, purchasing, "
        "manufacturing, packaging, labeling, or storage are qualified as approved suppliers in "
        "accordance with the supplier management process.")

    add_heading3(doc, "8.2.2 Competence, Awareness, and Training")
    add_body(doc, "Kytosan Bio establishes and maintains procedures for:")
    add_bullet(doc, "Identifying competency and training needs for all roles affecting product quality.")
    add_bullet(doc, "Providing required training and assessing its effectiveness proportionate to risk.")
    add_bullet(doc, "Making employees aware of the relevance of their activities to quality objectives.")
    add_bullet(doc, "Maintaining records of education, training, skills, and experience for all personnel.")
    add_body(doc,
        "Training on new or revised procedures is completed prior to implementation in the "
        "production or quality environment. Training effectiveness is assessed through methods "
        "proportionate to the risk associated with the subject matter.")

    add_heading2(doc, "8.3 Infrastructure")
    add_body(doc,
        "Kytosan Bio identifies, provides, and maintains the infrastructure necessary to ensure "
        "product conformance, including:")
    add_bullet(doc, "Facilities and workspaces, including cleanroom or controlled environment areas required for wound dressing manufacturing.")
    add_bullet(doc, "Manufacturing equipment, instruments, and tooling.")
    add_bullet(doc, "Computer hardware and software (validated per applicable requirements).")
    add_bullet(doc, "Utilities (e.g., purified water, compressed gases, HVAC).")
    add_bullet(doc, "Supporting services (e.g., sterilization contract services, calibration services).")
    add_body(doc,
        "Procedures are established for infrastructure maintenance, including required frequency "
        "of maintenance activities. Records of infrastructure maintenance are maintained.")

    add_heading2(doc, "8.4 Work Environment and Contamination Control")
    add_heading3(doc, "8.4.1 Work Environment")
    add_body(doc,
        "Kytosan Bio determines and manages the work environment to ensure product requirements are "
        "met. Requirements are established and maintained for:")
    add_bullet(doc, "Health, cleanliness, and clothing of personnel whose contact with the product or environment could affect product quality.")
    add_bullet(doc, "Monitoring and control of environmental conditions (temperature, humidity, particulates, bioburden) where these affect product quality.")
    add_bullet(doc, "Training or supervision for personnel working under special environmental conditions.")
    add_body(doc,
        "Given that the Kytosan Bio wound dressing is a sterile product, appropriate environmental "
        "controls are established and maintained to limit bioburden and particulate contamination "
        "during manufacturing and packaging processes prior to sterilization.")

    add_heading3(doc, "8.4.2 Contamination Control")
    add_body(doc,
        "Procedures are established and maintained for the control of contaminated or potentially "
        "contaminated product to prevent contamination of other product, the work environment, or "
        "personnel. Controls include:")
    add_bullet(doc, "Segregation of nonconforming or potentially contaminated materials.")
    add_bullet(doc, "Environmental monitoring program for manufacturing areas.")
    add_bullet(doc, "Personnel hygiene and gowning requirements.")
    add_bullet(doc, "Cleaning and disinfection procedures for manufacturing equipment and surfaces.")
    add_bullet(doc, "Bioburden monitoring of product prior to sterilization.")

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 9 — PRODUCT REALIZATION
    # ─────────────────────────────────────────────────────────────────────────
    add_heading1(doc, "9. Product Realization")

    add_heading2(doc, "9.1 Planning of Product Realization")
    add_body(doc,
        "Kytosan Bio establishes and maintains procedures for product realization planning to ensure "
        "that customer and product requirements are met. Product realization planning includes "
        "risk management and usability activities. Planning activities determine or develop:")
    add_bullet(doc, "Product quality objectives and requirements.")
    add_bullet(doc, "A product-specific quality plan.")
    add_bullet(doc, "Processes, documents, and resources specific to the product.")
    add_bullet(doc, "Product and process risks (risk management per ISO 14971).")
    add_bullet(doc, "Verification, validation, monitoring, inspection, and testing activities.")
    add_bullet(doc, "Product acceptance criteria.")
    add_body(doc,
        "Records of product realization planning activities are maintained. The product realization "
        "plan is reviewed and updated throughout the product lifecycle.")

    add_heading2(doc, "9.2 Customer-Related Processes")
    add_heading3(doc, "9.2.1 Determination of Requirements")
    add_body(doc,
        "Procedures are established to determine customer and product requirements, including "
        "requirements for delivery and post-delivery activities, applicable statutory and regulatory "
        "requirements, user training needs, and any other requirements necessary for the intended use "
        "of the product.")

    add_heading3(doc, "9.2.2 Review of Requirements")
    add_body(doc,
        "Requirements are reviewed prior to any commitment to supply product. The review ensures "
        "that requirements are defined and documented, regulatory requirements are met, and Kytosan "
        "Bio has the capability to meet defined requirements. Records of requirement reviews are "
        "maintained.")

    add_heading3(doc, "9.2.3 Communication")
    add_body(doc,
        "Effective communication arrangements are implemented with customers regarding product "
        "information, order handling, customer feedback and complaints, and field actions or "
        "regulatory notices. Kytosan Bio communicates with regulatory authorities as required by "
        "applicable regulations.")

    add_heading2(doc, "9.3 Design and Development Controls")
    add_body(doc,
        "Design and development controls are a critical element of the Kytosan Bio QMS, given the "
        "Class II medical device classification and the 510(k) pathway requirement. Design controls "
        "ensure that specified design requirements are met and that design outputs are suitable for "
        "manufacturing.")

    add_heading3(doc, "9.3.1 Design and Development Planning")
    add_body(doc,
        "Kytosan Bio implements a stage-gate design and development process for the systematic "
        "development of the Chitosan wound dressing and its associated manufacturing processes. "
        "Design plans are established, documented, and maintained that describe:")
    add_bullet(doc, "Design and development stages and gates.")
    add_bullet(doc, "Responsibilities and accountabilities for design activities.")
    add_bullet(doc, "Organizational and technical interfaces between functional areas.")
    add_bullet(doc, "Verification, validation, and design transfer milestones.")
    add_body(doc,
        "Design plans are reviewed and updated as the design evolves. Design outputs are traceable "
        "to design inputs.")

    add_heading3(doc, "9.3.2 Design Inputs")
    add_body(doc,
        "Design inputs are established to ensure that product requirements address:")
    add_bullet(doc, "Functional, performance, and safety requirements (e.g., antimicrobial efficacy, moisture vapor transmission rate, tensile strength).")
    add_bullet(doc, "Risk management outputs (ISO 14971).")
    add_bullet(doc, "Biocompatibility requirements (ISO 10993 series) — critical for Chitosan derived from shellfish chitin, requiring cytotoxicity, sensitization, and systemic toxicity testing.")
    add_bullet(doc, "Sterilization requirements and sterility assurance level (SAL) specification.")
    add_bullet(doc, "Applicable regulatory and labeling requirements.")
    add_bullet(doc, "Information derived from previous similar designs and market feedback.")
    add_body(doc,
        "Design inputs are reviewed and approved by designated individuals and are dated. "
        "A mechanism for addressing incomplete, ambiguous, or conflicting requirements is "
        "established. Records of design inputs are maintained in the Design History File.")

    add_heading3(doc, "9.3.3 Design Outputs")
    add_body(doc,
        "Design outputs are established in a form that enables verification against design inputs. "
        "Design outputs include:")
    add_bullet(doc, "Product specifications (chemical, physical, biological, and performance).")
    add_bullet(doc, "Manufacturing specifications and processes.")
    add_bullet(doc, "Sterilization method and process parameters.")
    add_bullet(doc, "Packaging and labeling specifications.")
    add_bullet(doc, "Acceptance criteria and test methods.")
    add_body(doc,
        "The Device Master Record, verification and validation results, technical documentation "
        "(EU Technical File), and the Declaration of Conformity (for CE Mark) comprise the design "
        "outputs. Design outputs are reviewed and approved prior to release.")

    add_heading3(doc, "9.3.4 Design Review")
    add_body(doc,
        "Formal, documented design reviews are planned and conducted at appropriate stages of "
        "development. Reviews include representatives of all functions concerned with the design "
        "stage and at least one independent reviewer. Records of design reviews and necessary "
        "actions are maintained.")

    add_heading3(doc, "9.3.5 Design Verification")
    add_body(doc,
        "Design verification is performed to confirm that design outputs meet design inputs. "
        "Verification activities for the Kytosan Bio wound dressing may include:")
    add_bullet(doc, "Antimicrobial efficacy testing (e.g., zone of inhibition, time-kill studies).")
    add_bullet(doc, "Biocompatibility testing per ISO 10993.")
    add_bullet(doc, "Physical and mechanical property testing (tensile strength, elongation, absorption).")
    add_bullet(doc, "Chemical characterization (degree of deacetylation, molecular weight, residuals).")
    add_bullet(doc, "Dimensional and visual inspection per specifications.")
    add_body(doc,
        "Verification plans include methods, acceptance criteria, and statistical rationale for "
        "sample sizes. Records are maintained in the DHF.")

    add_heading3(doc, "9.3.6 Design Validation")
    add_body(doc,
        "Design validation confirms that the product meets specified user needs and intended use. "
        "Validation is performed under defined conditions on representative initial production units "
        "or their equivalents, and includes:")
    add_bullet(doc, "Sterilization validation (including bioburden testing, sterility assurance level confirmation, dose setting, and dose audits) per applicable ISO 11135 or ISO 11137 standards.")
    add_bullet(doc, "Sterile barrier (packaging) validation per ISO 11607.")
    add_bullet(doc, "Shelf-life / accelerated aging validation.")
    add_bullet(doc, "Clinical evaluation in accordance with applicable regulatory requirements.")
    add_bullet(doc, "Usability / human factors evaluation.")
    add_body(doc,
        "Design validation is successfully completed prior to commercial distribution. "
        "Records of validation activities are maintained in the DHF.")

    add_heading3(doc, "9.3.7 Design Transfer")
    add_body(doc,
        "Design transfer activities ensure that design outputs are verified as suitable for "
        "manufacturing before becoming final production specifications. Transfer activities include "
        "process qualification (IQ/OQ/PQ) for critical manufacturing processes and sterilization "
        "processes. Records of design transfer activities are maintained.")

    add_heading3(doc, "9.3.8 Control of Design Changes")
    add_body(doc,
        "All design changes are reviewed, evaluated, and approved prior to implementation. "
        "Change reviews include evaluation of effects on constituent parts, product already "
        "distributed, manufacturing processes, affected risks, and any necessary regulatory updates "
        "(e.g., 510(k) submission or EU MDR technical documentation update). Records are maintained.")

    add_heading3(doc, "9.3.9 Design History File")
    add_body(doc,
        "A Design History File is established and maintained for the Kytosan Bio wound dressing. "
        "The DHF contains or references all records necessary to demonstrate that the design was "
        "developed in accordance with the approved design plan and QMS requirements, including all "
        "design inputs, outputs, reviews, verifications, validations, and transfer records.")

    add_heading2(doc, "9.4 Purchasing and Supplier Controls")
    add_heading3(doc, "9.4.1 Purchasing Process")
    add_body(doc,
        "Procedures are established to ensure that purchased products and services conform to "
        "specified requirements. Supplier qualification and evaluation criteria include:")
    add_bullet(doc, "Supplier ability to meet Kytosan Bio's quality and regulatory requirements.")
    add_bullet(doc, "Supplier performance history and quality system maturity.")
    add_bullet(doc, "Effect of the purchased product on Kytosan Bio product quality.")
    add_bullet(doc, "Risk associated with the supplied material or service.")
    add_body(doc,
        "Critical raw materials for the Chitosan wound dressing — including Chitosan polymer, "
        "carrier matrix polymers, sterilization service providers, and sterile packaging materials "
        "— receive heightened supplier controls commensurate with their risk. An Approved Vendor "
        "List (AVL) is maintained. Records of supplier evaluations are maintained.")

    add_heading3(doc, "9.4.2 Purchasing Information")
    add_body(doc,
        "Purchasing documents clearly describe or reference specified requirements including material "
        "description, quantity, specifications, acceptance requirements, and, where applicable, "
        "a requirement for the supplier to notify Kytosan Bio of changes to the purchased product "
        "prior to implementation.")

    add_heading3(doc, "9.4.3 Verification of Purchased Product")
    add_body(doc,
        "Incoming inspection procedures are established to verify that purchased product meets "
        "specified requirements, proportionate to supplier qualification status and risk. "
        "When requirements are verified at a supplier's premises, verification methods and "
        "product release approach are referenced in purchasing documents. Records of incoming "
        "inspection are maintained.")

    add_heading2(doc, "9.5 Production and Service Controls")
    add_heading3(doc, "9.5.1 Control of Production")
    add_body(doc,
        "Production activities are planned and executed under controlled conditions, including:")
    add_bullet(doc, "Product characteristics information and approved specifications.")
    add_bullet(doc, "Device History Records (DHRs) for each production lot.")
    add_bullet(doc, "Standard Operating Procedures and work instructions.")
    add_bullet(doc, "Appropriate manufacturing equipment in qualified condition.")
    add_bullet(doc, "In-process monitoring and measurement as required.")
    add_bullet(doc, "Acceptance activities, release, and distribution controls.")

    add_heading3(doc, "9.5.2 Cleanliness of Product")
    add_body(doc,
        "Procedures are established for product cleanliness because the Kytosan Bio wound dressing "
        "cannot be cleaned post-assembly/pre-sterilization, and cleanliness is critical to product "
        "quality and sterility assurance. Bioburden limits are specified and tested prior to "
        "sterilization.")

    add_heading3(doc, "9.5.3 Sterile Medical Device Requirements")
    add_body(doc,
        "Records of process parameters for the sterilization process used for each production lot "
        "are maintained and are fully traceable to each production lot. Sterilization process "
        "parameters are defined in the DMR and controlled to ensure consistency with the validated "
        "sterilization process.")

    add_heading3(doc, "9.5.4 Validation of Production Processes")
    add_body(doc,
        "Processes where subsequent monitoring and measurement cannot fully verify results are "
        "validated. For Kytosan Bio, this includes sterilization processes, sterile barrier "
        "sealing processes, and any other special processes. Validation activities include:")
    add_bullet(doc, "Installation Qualification (IQ).")
    add_bullet(doc, "Operational Qualification (OQ).")
    add_bullet(doc, "Performance Qualification (PQ).")
    add_body(doc,
        "A Master Validation Plan (MVP) identifies all production processes requiring initial "
        "validation and revalidation criteria and current status. Records of validation activities, "
        "results, and conclusions are maintained.")

    add_heading3(doc, "9.5.5 Validation of Sterilization and Sterile Barrier Systems")
    add_body(doc,
        "Sterilization processes and sterile barrier systems are validated prior to initial use "
        "and following product or process changes per applicable standards (ISO 11135, ISO 11137, "
        "ISO 11607). Records of sterilization process validation are maintained.")

    add_heading3(doc, "9.5.6 Identification and Traceability")
    add_body(doc,
        "Procedures are established for product identification throughout all stages of realization "
        "and distribution. Lot numbers are used to identify each unit or lot of finished product. "
        "Records of lot numbers are maintained in DHRs. If required by regulation, a Unique Device "
        "Identification (UDI) system is assigned in accordance with applicable FDA UDI regulations.")

    add_heading3(doc, "9.5.7 Preservation of Product")
    add_body(doc,
        "Procedures are established for the identification, handling, packaging, storage, "
        "preservation, and delivery of finished product to maintain product conformity. "
        "Special storage conditions (e.g., temperature, humidity) are documented and controlled. "
        "Shelf-life requirements for the wound dressing are specified and validated.")

    add_heading2(doc, "9.6 Control of Monitoring and Measuring Equipment")
    add_body(doc,
        "Procedures are established to identify, control, calibrate, and maintain all equipment "
        "used to test and monitor product and processes during design validation and production. "
        "Key controls include:")
    add_bullet(doc, "Equipment is verified to ensure calibration accuracy and precision required for its intended measurement.")
    add_bullet(doc, "Calibration standards are traceable to the National Institute of Standards and Technology (NIST) or other recognized international standards.")
    add_bullet(doc, "Equipment is marked with its calibration status and is safeguarded from unauthorized adjustments.")
    add_bullet(doc, "Equipment found out of calibration is removed from service, quarantined, and assessed for risk to prior measurements.")
    add_bullet(doc, "Records of equipment calibration and verification are maintained.")

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 10 — MEASUREMENT, ANALYSIS, AND IMPROVEMENT
    # ─────────────────────────────────────────────────────────────────────────
    add_heading1(doc, "10. Measurement, Analysis, and Improvement")

    add_heading2(doc, "10.1 General")
    add_body(doc,
        "Kytosan Bio establishes and maintains procedures for monitoring, measurement, analysis, "
        "and improvement processes to demonstrate product conformity, ensure QMS conformity, "
        "and maintain QMS effectiveness. Statistical techniques are identified and applied as "
        "appropriate to support data analysis activities.")

    add_heading2(doc, "10.2 Monitoring and Measurement")
    add_heading3(doc, "10.2.1 Customer Feedback and Satisfaction")
    add_body(doc,
        "Procedures are established for monitoring customer perception and whether product "
        "requirements have been met. Data sources include customer complaints, surveys, market "
        "feedback, and post-market surveillance. Information serves as input into risk management "
        "and product improvement processes.")

    add_heading3(doc, "10.2.2 Complaint Handling")
    add_body(doc,
        "Procedures are established for timely complaint handling, including:")
    add_bullet(doc, "Receiving and recording all feedback.")
    add_bullet(doc, "Determining whether feedback qualifies as a reportable complaint.")
    add_bullet(doc, "Investigating complaints to determine root cause.")
    add_bullet(doc, "Determining Medical Device Reporting (MDR) obligations (FDA) or vigilance reporting obligations (EU MDR).")
    add_bullet(doc, "Initiating CAPA as appropriate.")
    add_body(doc, "Records of complaint handling activities are maintained.")

    add_heading3(doc, "10.2.3 Reporting to Regulatory Authorities")
    add_body(doc,
        "Procedures are established for providing mandatory reports to regulatory authorities for "
        "complaints meeting specified adverse event reporting criteria, including FDA MDR reports "
        "(21 CFR Part 803) and EU MDR vigilance reports. Records of regulatory reporting are maintained.")

    add_heading3(doc, "10.2.4 Internal Audit")
    add_body(doc,
        "Procedures are established for conducting internal audits to assess the effectiveness of "
        "the QMS and compliance with applicable standards and regulations. Internal audits of the "
        "entire QMS are conducted at least annually. Audits are performed by trained Lead Auditors "
        "independent of the areas being audited. Audit findings and corrective action responses are "
        "documented and presented to management. Follow-up activities verify implementation and "
        "effectiveness of corrective actions. Records of internal audits are maintained.")

    add_heading3(doc, "10.2.5 Monitoring and Measurement of Processes")
    add_body(doc,
        "Procedures are established for monitoring and/or measurement of QMS processes. "
        "Process metrics demonstrate the ability of processes to achieve specified requirements. "
        "When process requirements are not achieved, corrective action is taken as appropriate.")

    add_heading3(doc, "10.2.6 Monitoring and Measurement of Product")
    add_body(doc,
        "Procedures are established for monitoring and measuring product characteristics at "
        "appropriate stages of the product realization process. Finished product is not released "
        "for distribution until all DMR activities are satisfactorily completed and authorized. "
        "Records of acceptance activities, including the identity of the individuals performing "
        "acceptance activities, are maintained.")

    add_heading2(doc, "10.3 Control of Nonconforming Product")
    add_heading3(doc, "10.3.1 General")
    add_body(doc,
        "Procedures are established to ensure that product not conforming to requirements is "
        "identified, segregated, and prevented from unintended use or distribution. "
        "Responsibilities and authorities for nonconforming material identification, "
        "documentation, evaluation, and disposition are defined.")

    add_heading3(doc, "10.3.2 Pre-Delivery Nonconformances")
    add_body(doc,
        "Nonconforming product detected prior to delivery is managed by one of the following "
        "dispositions: rework, rejection/scrap, or acceptance by concession (use-as-is). "
        "Concessions require documented justification and approval. Records are maintained.")

    add_heading3(doc, "10.3.3 Post-Delivery Nonconformances")
    add_body(doc,
        "When nonconforming product is detected after or during distribution, appropriate "
        "action is taken proportionate to the effects or potential effects of the nonconformity. "
        "Procedures are established for issuing advisory notices and field safety corrective "
        "actions (FSCAs) in accordance with applicable regulatory requirements (21 CFR Part 806; "
        "EU MDR Article 83-89).")

    add_heading3(doc, "10.3.4 Rework")
    add_body(doc,
        "Rework of nonconforming product is performed according to documented, approved rework "
        "instructions that have undergone the same review and approval process as the original "
        "procedure. Completed reworked product is verified against acceptance criteria before "
        "release. Records of rework are maintained.")

    add_heading2(doc, "10.4 Analysis of Data")
    add_body(doc,
        "Procedures are established to collect and analyze appropriate data to demonstrate the "
        "suitability and effectiveness of the QMS. Data analysis provides information on:")
    add_bullet(doc, "Customer feedback and satisfaction trends.")
    add_bullet(doc, "Product conformity to requirements.")
    add_bullet(doc, "Process and product characteristics and trends.")
    add_bullet(doc, "Supplier performance.")
    add_bullet(doc, "Internal audit results and trends.")
    add_body(doc,
        "Statistical techniques are applied as appropriate. If analysis shows the QMS is not "
        "suitable, adequate, or effective, the results are used as inputs to improvement activities. "
        "Records of data analysis results are maintained.")

    add_heading2(doc, "10.5 Improvement — CAPA and Preventive Action")
    add_heading3(doc, "10.5.1 General")
    add_body(doc,
        "Kytosan Bio identifies and implements changes necessary to maintain QMS effectiveness, "
        "suitability, and adequacy through the use of the Quality Policy, quality objectives, "
        "audit results, data analysis, post-market surveillance, corrective actions, preventive "
        "actions, and management review.")

    add_heading3(doc, "10.5.2 Corrective Action (CA)")
    add_body(doc,
        "Procedures are established to eliminate the causes of nonconformities to prevent "
        "recurrence. Corrective actions are proportionate to the problems and risks. "
        "The CAPA procedure defines requirements for:")
    add_bullet(doc, "Reviewing nonconformities (including complaints).")
    add_bullet(doc, "Investigating root causes.")
    add_bullet(doc, "Evaluating whether corrective action is needed to prevent recurrence.")
    add_bullet(doc, "Planning, documenting, and implementing corrective action.")
    add_bullet(doc, "Verifying that corrective action does not adversely affect regulatory compliance or product safety.")
    add_bullet(doc, "Reviewing the effectiveness of corrective actions taken.")
    add_body(doc, "Records of investigations and corrective actions are maintained.")

    add_heading3(doc, "10.5.3 Preventive Action (PA)")
    add_body(doc,
        "Procedures are established to eliminate the causes of potential nonconformities to "
        "prevent their occurrence. Preventive actions are proportionate to potential risks. "
        "The process includes identification of potential nonconformities, root cause analysis, "
        "implementation of preventive actions, and review of effectiveness. Records are maintained.")

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 11 — REGULATORY COMPLIANCE
    # ─────────────────────────────────────────────────────────────────────────
    add_heading1(doc, "11. Regulatory Compliance")

    add_heading2(doc, "11.1 FDA 21 CFR Part 820 Compliance Summary")
    add_body(doc,
        "The table below summarizes the relationship between FDA 21 CFR Part 820 requirements "
        "and the Kytosan Bio QMS procedures. Document numbers are placeholders pending full QMS "
        "build-out.")

    fda_data = [
        ("§820.20", "Management Responsibility",       "QM-001 (this document)"),
        ("§820.22", "Quality Audit",                   "SOP-QA-001: Internal Audit Program"),
        ("§820.25", "Personnel",                       "SOP-HR-001: Training and Competency"),
        ("§820.30", "Design Controls",                 "SOP-DD-001: Design and Development Controls"),
        ("§820.40", "Document Controls",               "SOP-DC-001: Document Change Control"),
        ("§820.50", "Purchasing Controls",             "SOP-SC-001: Supplier and Purchasing Management"),
        ("§820.60", "Identification",                  "SOP-OP-001: Identification and Traceability"),
        ("§820.65", "Traceability",                    "SOP-OP-001: Identification and Traceability"),
        ("§820.70", "Production and Process Controls", "SOP-MFG-001: Manufacturing Process Controls"),
        ("§820.72", "Inspection, Measuring, and Test Equipment", "SOP-MFG-002: Equipment Calibration and Maintenance"),
        ("§820.75", "Process Validation",              "SOP-VAL-001: Process Validation"),
        ("§820.80", "Receiving, In-Process, and Finished Device Acceptance", "SOP-QC-001: Acceptance Activities"),
        ("§820.86", "Acceptance Status",               "SOP-QC-001: Acceptance Activities"),
        ("§820.90", "Nonconforming Product",           "SOP-QA-002: Nonconforming Material Control"),
        ("§820.100","CAPA",                            "SOP-QA-003: Corrective and Preventive Action"),
        ("§820.120","Device Labeling",                 "SOP-OP-002: Labeling and Packaging Controls"),
        ("§820.130","Device Packaging",                "SOP-OP-002: Labeling and Packaging Controls"),
        ("§820.140","Handling",                        "SOP-MFG-001: Manufacturing Process Controls"),
        ("§820.150","Storage",                         "SOP-OP-003: Storage and Distribution"),
        ("§820.160","Distribution",                    "SOP-OP-003: Storage and Distribution"),
        ("§820.180","General Record Requirements",     "SOP-DC-002: Record Control"),
        ("§820.181","Device Master Record",            "QM-001; Product DMR"),
        ("§820.184","Device History Record",           "SOP-QC-001: Acceptance Activities"),
        ("§820.186","Quality System Record",           "QM-001 (this document)"),
        ("§820.198","Complaint Files",                 "SOP-PMS-001: Complaint Handling and Reporting"),
    ]

    fda_tbl = doc.add_table(rows=len(fda_data) + 1, cols=3)
    fda_tbl.style = "Table Grid"
    add_table_header_row(fda_tbl, ["21 CFR §820 Clause", "Requirement", "Kytosan Bio Procedure"])
    for i, row_d in enumerate(fda_data):
        add_table_data_row(fda_tbl, i + 1, row_d, shaded=(i % 2 == 0))
    set_cell_borders(fda_tbl)
    doc.add_paragraph()

    add_heading2(doc, "11.2 EU MDR 2017/745 Requirements — Class IIa, CE Mark")
    add_body(doc,
        "Kytosan Bio's antimicrobial wound dressing is classified as a Class IIa medical device "
        "under EU MDR 2017/745 (Rule 7: Devices that are intended to be used in direct contact "
        "with wounds). CE Mark certification requires conformity assessment by a Notified Body, "
        "Technical Documentation, and a Quality Management System audited against Annex IX "
        "(QMS + Technical Documentation review) or Annex XI (Product Verification).")

    add_body(doc, "Key EU MDR requirements for Kytosan Bio include:", bold=True)
    add_bullet(doc, "Technical Documentation (Annex II and III) — complete product description, design and manufacturing information, General Safety and Performance Requirements (GSPR) conformity, risk management, clinical evaluation, performance data.")
    add_bullet(doc, "Clinical Evaluation (Article 61 and Annex XIV) — clinical data demonstrating safety and performance for the intended purpose.")
    add_bullet(doc, "Post-Market Clinical Follow-Up (PMCF) Plan and Report.")
    add_bullet(doc, "Post-Market Surveillance (PMS) Plan and Periodic Safety Update Report (PSUR) — required at least every 2 years for Class IIa devices.")
    add_bullet(doc, "Summary of Safety and Clinical Performance (SSCP) — publicly accessible summary for Class IIa and above (Article 32).")
    add_bullet(doc, "Unique Device Identification (UDI) registration in the EUDAMED database.")
    add_bullet(doc, "Declaration of Conformity (DoC) signed by the Authorized Representative.")
    add_bullet(doc, "Authorized Representative (AR) designation in the EU (if Kytosan Bio is based outside the EU).")
    add_bullet(doc, "Importer responsibilities (Article 13) if applicable.")

    add_heading2(doc, "11.3 ISO 13485:2016 Cross-Reference Matrix")

    iso_data = [
        ("4.1", "General QMS Requirements",             "QM-001, Section 6.1; SOP-DC-001"),
        ("4.2", "Documentation Requirements",           "QM-001, Section 6.2-6.4; SOP-DC-001; SOP-DC-002"),
        ("5.1", "Management Commitment",                "QM-001, Section 7.1"),
        ("5.2", "Customer Focus",                       "QM-001, Section 7.2; SOP-PMS-001"),
        ("5.3", "Quality Policy",                       "QM-001, Section 2 & 7.3; Appendix A"),
        ("5.4", "Planning",                             "QM-001, Section 7.4"),
        ("5.5", "Responsibility, Authority, Communication", "QM-001, Section 7.5"),
        ("5.6", "Management Review",                    "QM-001, Section 7.6; SOP-QA-004"),
        ("6.1", "Provision of Resources",               "QM-001, Section 8.1"),
        ("6.2", "Human Resources",                      "QM-001, Section 8.2; SOP-HR-001"),
        ("6.3", "Infrastructure",                       "QM-001, Section 8.3; SOP-MFG-002"),
        ("6.4", "Work Environment",                     "QM-001, Section 8.4; SOP-MFG-003"),
        ("7.1", "Planning of Product Realization",      "QM-001, Section 9.1; SOP-DD-001"),
        ("7.2", "Customer-Related Processes",           "QM-001, Section 9.2; SOP-PMS-001"),
        ("7.3", "Design and Development",               "QM-001, Section 9.3; SOP-DD-001"),
        ("7.4", "Purchasing",                           "QM-001, Section 9.4; SOP-SC-001"),
        ("7.5", "Production and Service Provision",     "QM-001, Section 9.5; SOP-MFG-001; SOP-VAL-001"),
        ("7.6", "Control of Monitoring and Measuring Equipment", "QM-001, Section 9.6; SOP-MFG-002"),
        ("8.1", "General — Measurement, Analysis, Improvement", "QM-001, Section 10.1"),
        ("8.2", "Monitoring and Measurement",           "QM-001, Section 10.2; SOP-QA-001; SOP-PMS-001"),
        ("8.3", "Control of Nonconforming Product",     "QM-001, Section 10.3; SOP-QA-002"),
        ("8.4", "Analysis of Data",                     "QM-001, Section 10.4"),
        ("8.5", "Improvement",                          "QM-001, Section 10.5; SOP-QA-003"),
    ]

    iso_tbl = doc.add_table(rows=len(iso_data) + 1, cols=3)
    iso_tbl.style = "Table Grid"
    add_table_header_row(iso_tbl, ["ISO 13485:2016 Clause", "Requirement", "Kytosan Bio Reference"])
    for i, row_d in enumerate(iso_data):
        add_table_data_row(iso_tbl, i + 1, row_d, shaded=(i % 2 == 0))
    set_cell_borders(iso_tbl)
    doc.add_paragraph()

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 12 — DOCUMENT CONTROL LIST
    # ─────────────────────────────────────────────────────────────────────────
    add_heading1(doc, "12. Document Control List — Referenced QMS Procedures")
    add_body(doc,
        "The following procedures and documents are referenced by this Quality Manual. "
        "All documents are controlled within the Kytosan Bio electronic QMS (eQMS). "
        "Document numbers and revisions are subject to change as the QMS matures.")

    doc_list = [
        ("QM-001",       "Quality Manual",                                             "This document"),
        ("SOP-DC-001",   "Document Change Control",                                    "Document Control"),
        ("SOP-DC-002",   "Record Control and Retention",                               "Document Control"),
        ("SOP-DC-003",   "Document Format and Numbering",                              "Document Control"),
        ("SOP-DC-004",   "Good Documentation Practices",                               "Document Control"),
        ("SOP-HR-001",   "Training, Competency, and Onboarding",                       "Human Resources"),
        ("SOP-QA-001",   "Internal Audit Program",                                     "Quality Assurance"),
        ("SOP-QA-002",   "Nonconforming Material Control",                             "Quality Assurance"),
        ("SOP-QA-003",   "Corrective and Preventive Action (CAPA)",                   "Quality Assurance"),
        ("SOP-QA-004",   "Management Review of the Quality System",                   "Quality Assurance"),
        ("SOP-QA-005",   "Statistical Techniques",                                     "Quality Assurance"),
        ("SOP-DD-001",   "Design and Development Controls",                            "R&D / Design"),
        ("SOP-DD-002",   "Risk Management (ISO 14971)",                                "R&D / Design"),
        ("SOP-DD-003",   "Biocompatibility Evaluation (ISO 10993)",                   "R&D / Design"),
        ("SOP-DD-004",   "Clinical Evaluation and Post-Market Clinical Follow-Up",    "R&D / Regulatory"),
        ("SOP-SC-001",   "Supplier Qualification and Purchasing Management",           "Supply Chain"),
        ("SOP-SC-002",   "Approved Vendor List Management",                            "Supply Chain"),
        ("SOP-MFG-001",  "Manufacturing Process Controls",                             "Manufacturing"),
        ("SOP-MFG-002",  "Equipment Calibration and Preventive Maintenance",          "Manufacturing"),
        ("SOP-MFG-003",  "Environmental Controls and Monitoring",                     "Manufacturing"),
        ("SOP-MFG-004",  "Manufacturing Cleaning and Waste Control",                  "Manufacturing"),
        ("SOP-VAL-001",  "Process Validation (IQ/OQ/PQ)",                             "Validation"),
        ("SOP-VAL-002",  "Sterilization Validation",                                  "Validation"),
        ("SOP-VAL-003",  "Packaging / Sterile Barrier Validation (ISO 11607)",        "Validation"),
        ("SOP-VAL-004",  "Software Validation",                                       "Validation"),
        ("SOP-QC-001",   "Incoming Inspection and Raw Material Testing",               "Quality Control"),
        ("SOP-QC-002",   "In-Process Inspection and Testing",                         "Quality Control"),
        ("SOP-QC-003",   "Final Acceptance Activities and Finished Product Release",  "Quality Control"),
        ("SOP-QC-004",   "Acceptance Status and Acceptance Activities",               "Quality Control"),
        ("SOP-OP-001",   "Identification, Traceability, and Lot Number Assignment",   "Operations"),
        ("SOP-OP-002",   "Labeling and Packaging Controls",                           "Operations"),
        ("SOP-OP-003",   "Storage, Handling, and Distribution",                       "Operations"),
        ("SOP-OP-004",   "Expiration Date and Shelf-Life Assignment",                 "Operations"),
        ("SOP-PMS-001",  "Customer Feedback, Complaint Handling, and Reporting",      "Post-Market"),
        ("SOP-PMS-002",  "Medical Device Reporting (FDA MDR — 21 CFR Part 803)",      "Post-Market"),
        ("SOP-PMS-003",  "EU MDR Vigilance Reporting",                                "Post-Market"),
        ("SOP-PMS-004",  "Corrections and Removals (21 CFR Part 806)",                "Post-Market"),
        ("SOP-PMS-005",  "Post-Market Surveillance Plan and PSUR",                   "Post-Market"),
        ("SOP-REG-001",  "510(k) Premarket Notification Management",                 "Regulatory"),
        ("SOP-REG-002",  "CE Mark and EU Technical Documentation Management",         "Regulatory"),
        ("SOP-REG-003",  "UDI Registration (FDA GUDID and EUDAMED)",                 "Regulatory"),
    ]

    doc_tbl = doc.add_table(rows=len(doc_list) + 1, cols=3)
    doc_tbl.style = "Table Grid"
    add_table_header_row(doc_tbl, ["Document Number", "Document Title", "Function"])
    for i, row_d in enumerate(doc_list):
        add_table_data_row(doc_tbl, i + 1, row_d, shaded=(i % 2 == 0))
    set_cell_borders(doc_tbl)
    doc.add_paragraph()

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # APPENDIX A — QUALITY POLICY STATEMENT
    # ─────────────────────────────────────────────────────────────────────────
    add_heading1(doc, "Appendix A — Quality Policy Statement")

    add_body(doc, "KYTOSAN BIO — QUALITY POLICY", bold=True)
    add_body(doc, "Document Number: QM-001.02  |  Revision: 0.1 DRAFT  |  Effective Date: May 4, 2026")
    doc.add_paragraph()

    p_qp = doc.add_paragraph()
    p_qp.paragraph_format.left_indent = Inches(0.5)
    p_qp.paragraph_format.right_indent = Inches(0.5)
    rqp = p_qp.add_run(
        "Kytosan Bio is dedicated to improving patient outcomes through the development and supply "
        "of safe, effective, and innovative Chitosan-based wound care solutions. We are committed to "
        "continuously improving our products and processes by maintaining rigorous standards of "
        "product quality from raw material receipt through finished product release; ensuring full "
        "compliance with applicable regulatory requirements including FDA 21 CFR Part 820, "
        "ISO 13485:2016, and EU MDR 2017/745; and sustaining an effective Quality Management System "
        "that is regularly reviewed for suitability, adequacy, and compliance.\n\n"
        "All Kytosan Bio employees share responsibility for quality and are empowered to identify "
        "and escalate quality issues. Our commitment to quality is inseparable from our commitment "
        "to patient safety."
    )
    rqp.font.name = "Calibri"
    rqp.font.size = Pt(11)
    rqp.font.italic = True
    rqp.font.color.rgb = NAVY

    doc.add_paragraph()
    add_body(doc, "Signed and approved by:", bold=True)
    doc.add_paragraph()

    sig_tbl = doc.add_table(rows=3, cols=3)
    sig_tbl.style = "Table Grid"
    add_table_header_row(sig_tbl, ["Name", "Title", "Signature / Date"])
    sig_data = [
        ("[NAME]", "Chief Executive Officer",                          "[SIGNATURE]  [DATE]"),
        ("[NAME]", "Head of Quality and Regulatory Affairs (Mgmt Rep)", "[SIGNATURE]  [DATE]"),
    ]
    for i, row_d in enumerate(sig_data):
        add_table_data_row(sig_tbl, i + 1, row_d, shaded=False)
    set_cell_borders(sig_tbl)
    doc.add_paragraph()

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # APPENDIX B — ORGANIZATIONAL STRUCTURE
    # ─────────────────────────────────────────────────────────────────────────
    add_heading1(doc, "Appendix B — Organizational Structure")
    add_body(doc,
        "The following describes the Kytosan Bio organizational structure. An updated organization "
        "chart is maintained in the eQMS and is reviewed and updated at each Management Review.")

    add_body(doc, "Kytosan Bio Reporting Structure (Placeholder — Early Stage):", bold=True)
    doc.add_paragraph()

    org_lines = [
        ("CEO / Founder", 0),
        ("Head of Quality and Regulatory Affairs (Management Representative)", 1),
        ("Quality Assurance / QMS Administrator", 2),
        ("Regulatory Affairs Specialist", 2),
        ("Head of Research and Development", 1),
        ("Senior Scientist — Materials / Chitosan Chemistry", 2),
        ("R&D Associate", 2),
        ("Clinical / Biocompatibility Specialist (Contract)", 2),
        ("Head of Operations / Manufacturing", 1),
        ("Manufacturing Associate(s)", 2),
        ("Supply Chain / Purchasing Coordinator", 2),
        ("Administration / Finance", 1),
    ]

    for title, level in org_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5 * level)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(("    " * level) + ("└── " if level > 0 else "") + title)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.bold = (level == 0)
        run.font.color.rgb = NAVY if level == 0 else (DARK_BLUE if level == 1 else BLACK)

    doc.add_paragraph()
    add_body(doc,
        "Note: As an early-stage company, Kytosan Bio currently operates with a small team. "
        "Individuals may hold multiple roles consistent with the requirements of FDA 21 CFR Part 820 "
        "and ISO 13485:2016. Role definitions and reporting lines will expand as the company grows. "
        "An Authorized Representative in the EU must be designated prior to CE Mark application "
        "if Kytosan Bio is headquartered outside the European Union.", italic=True)

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # APPENDIX C — PROCESS MAP / TURTLE DIAGRAM
    # ─────────────────────────────────────────────────────────────────────────
    add_heading1(doc, "Appendix C — Process Map / Turtle Diagram")
    add_body(doc,
        "The Kytosan Bio QMS is structured around a process-based approach consistent with ISO "
        "13485:2016 and FDA 21 CFR Part 820. The process interactions are described below. "
        "A visual process map or turtle diagram is maintained in the eQMS.")

    add_heading2(doc, "C.1 High-Level Process Flow")
    add_body(doc, "The Kytosan Bio QMS encompasses the following major process categories:")

    proc_data = [
        ("Management Processes",
         "Quality Policy; Quality Objectives; Management Review; Resource Management; Internal Audit",
         "Strategic direction; QMS performance data; Resource allocation decisions"),
        ("Design and Development",
         "Design planning; Design inputs/outputs; Verification; Validation; Risk Management; Clinical Evaluation; Design Transfer",
         "DHF; DMR; Validated product; Technical Documentation; 510(k) / CE Mark submission"),
        ("Supplier and Purchasing",
         "Supplier qualification; Purchasing; Incoming inspection",
         "Approved Vendor List; Qualified raw materials and components"),
        ("Manufacturing / Product Realization",
         "Manufacturing execution; In-process controls; Sterilization; Packaging; Labeling; Final release",
         "Finished sterile wound dressings; DHRs; Release records"),
        ("Measurement, Analysis, Improvement",
         "Internal audits; Complaint handling; CAPA; Data analysis; Post-market surveillance",
         "Audit reports; CAPA records; PMS data; Continuous improvement actions"),
        ("Regulatory and Post-Market",
         "510(k) management; CE Mark Technical Documentation; MDR/Vigilance reporting; UDI; PSUR",
         "Marketing authorizations; Regulatory submissions; Post-market reports"),
    ]
    proc_tbl = doc.add_table(rows=len(proc_data) + 1, cols=3)
    proc_tbl.style = "Table Grid"
    add_table_header_row(proc_tbl, ["Process Category", "Key Sub-Processes", "Key Outputs"])
    for i, row_d in enumerate(proc_data):
        add_table_data_row(proc_tbl, i + 1, row_d, shaded=(i % 2 == 0))
    set_cell_borders(proc_tbl)
    doc.add_paragraph()

    add_heading2(doc, "C.2 Turtle Diagram — Manufacturing Process (Example)")
    add_body(doc,
        "A turtle diagram for the Kytosan Bio wound dressing manufacturing process is described below. "
        "A visual diagram is maintained in the eQMS.")

    turtle_tbl = doc.add_table(rows=7, cols=2)
    turtle_tbl.style = "Table Grid"
    add_table_header_row(turtle_tbl, ["Turtle Diagram Element", "Content for Manufacturing Process"])
    turtle_data = [
        ("What — Inputs",
         "Raw materials (Chitosan, polymer matrix components); Packaging materials; "
         "Process parameters from DMR; Approved work instructions; Calibrated equipment"),
        ("What — Outputs",
         "Sterile, labeled, packaged antimicrobial wound dressings; DHR; Release records; "
         "Traceability records"),
        ("How — Process Steps",
         "Receiving inspection → Material preparation → Manufacturing / forming → "
         "In-process inspection → Packaging and labeling → Sterilization → "
         "Post-sterilization inspection → Final release"),
        ("Who — Competencies",
         "Trained manufacturing associates; Qualified sterilization operators; "
         "QC personnel; Authorized release personnel"),
        ("With What — Equipment / Methods",
         "Formulation / mixing equipment; Casting / forming equipment; "
         "Cleanroom environment; Sterilization equipment (or contract sterilizer); "
         "Calibrated measurement and test equipment"),
        ("How Do We Know It Is Working — KPIs / Metrics",
         "Yield per production lot; In-process nonconformance rate; "
         "Sterility test pass rate; On-time release rate; Complaint rate; "
         "Calibration compliance rate"),
    ]
    for i, row_d in enumerate(turtle_data):
        add_table_data_row(turtle_tbl, i + 1, row_d, shaded=(i % 2 == 0))
    set_cell_borders(turtle_tbl)
    doc.add_paragraph()

    add_horizontal_rule(doc)
    add_body(doc, "— End of Kytosan Bio Quality Manual QM-001 Rev 0.1 DRAFT —", bold=True)
    p_end = doc.paragraphs[-1]
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_body(doc,
        "CONFIDENTIAL — This document is the property of Kytosan Bio and contains proprietary "
        "information. Unauthorized disclosure, reproduction, or distribution is prohibited. "
        "This is a controlled document. Printed copies are uncontrolled.", italic=True)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ─────────────────────────────────────────────────────────────────────────
    # SAVE
    # ─────────────────────────────────────────────────────────────────────────
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Saved: {OUTPUT_PATH}")
    print(f"File size: {OUTPUT_PATH.stat().st_size:,} bytes")


if __name__ == "__main__":
    build_document()
